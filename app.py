import customtkinter as ctk
from gpt4all import GPT4All
import os
from dotenv import load_dotenv
import tkinter.messagebox as messagebox
from tkinter import filedialog
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import Inches
from docx2pdf import convert
import tkinter
import json
from datetime import datetime

# Charger les variables d'environnement
load_dotenv()

class LetterGeneratorApp:
    def __init__(self):
        self.window = ctk.CTk()
        self.window.title("Générateur de Lettre de Motivation")
        self.window.geometry("1200x800")
        
        # Frame principale avec défilement
        self.main_frame = ctk.CTkScrollableFrame(self.window)
        self.main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Frame pour le message de statut (créé en premier)
        self.status_frame = ctk.CTkFrame(self.window)
        self.status_frame.pack(fill="x", padx=20, pady=5, side="bottom")
        
        # Label pour afficher les messages de statut
        self.status_label = ctk.CTkLabel(
            self.status_frame,
            text="",
            font=("Arial", 12),
            text_color="#28a745"  # Vert pour les messages de succès
        )
        self.status_label.pack(pady=5)
        
        # Initialiser les styles par défaut
        self.template_styles = [{
            "font_name": "Times New Roman",
            "font_size": 11,
            "bold": False,
            "alignment": 0,  # Gauche
            "line_spacing": 1.15,
            "space_before": Pt(0),
            "space_after": Pt(8),
            "first_line_indent": None
        }]
        
        # Initialiser le modèle
        self.llm = None
        
        # Créer un dossier pour sauvegarder les données si nécessaire
        self.save_dir = os.path.join(os.path.expanduser("~"), ".lettre_motivation_ai")
        os.makedirs(self.save_dir, exist_ok=True)
        self.save_file = os.path.join(self.save_dir, "last_session.json")
        self.templates_file = os.path.join(self.save_dir, "custom_templates.json")
        
        # Chemin de sauvegarde par défaut pour les fichiers Word
        self.word_save_path = os.path.expanduser("~")
        
        # Initialiser le dictionnaire des modèles
        self.custom_templates = {}
        
        # Charger les modèles personnalisés
        self.load_custom_templates()
        
        # Stocker les informations de mise en page
        self.template_style = {
            "font_name": "Times New Roman",
            "font_size": 11,
            "alignment": "left",
            "line_spacing": 1.15
        }
        
        # Définir les couleurs des marqueurs
        self.marker_colors = {
            "company": "#34C759",     # Vert clair pour l'entreprise
            "position": "#3498DB",    # Bleu ciel pour le poste
            "duration": "#9B59B6",    # Violet clair pour la durée
            "start_date": "#F7DC6F",  # Jaune clair pour la date de début
            "today_date": "#FFC5C5",  # Rose clair pour la date du jour
            "custom": "#FFA07A"       # Orange clair pour le paragraphe personnalisé
        }
        
        # Créer l'interface
        self.create_widgets()
        
        # Charger les dernières données
        self.load_last_session()
        
        # Sauvegarder à la fermeture
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # Charger le modèle au démarrage
        self.load_model()
        
        # Configurer la sauvegarde automatique
        self.setup_auto_save()
    
    def create_widgets(self):
        # === PREMIÈRE PARTIE : OFFRE D'EMPLOI ===
        self.create_input_section(self.main_frame)
        
        # Séparateur
        separator = ctk.CTkFrame(self.main_frame, height=2)
        separator.pack(fill="x", pady=20)
        
        # === DEUXIÈME PARTIE : MODÈLE DE LETTRE ===
        self.create_template_section(self.main_frame)
        
        # Section résultat
        result_frame = ctk.CTkFrame(self.main_frame)
        result_frame.pack(fill="x", padx=20, pady=10)
        
        # En-tête avec titre et bouton d'export
        header_frame = ctk.CTkFrame(result_frame)
        header_frame.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(header_frame, text="Lettre générée", font=("Arial", 16, "bold")).pack(side="left", pady=5)
        
        # Frame pour les boutons d'action
        button_frame = ctk.CTkFrame(result_frame)
        button_frame.pack(fill="x", padx=20, pady=5)

        # Première ligne : nom du fichier et dossier
        filename_frame = ctk.CTkFrame(button_frame)
        filename_frame.pack(fill="x", padx=5, pady=(0, 5))

        # Label et champ pour le nom du fichier
        filename_label = ctk.CTkLabel(filename_frame, text="Nom du fichier:")
        filename_label.pack(side="left", padx=5)
        
        self.filename_entry = ctk.CTkEntry(filename_frame, width=200)
        self.filename_entry.pack(side="left", padx=5, fill="x", expand=True)

        # Bouton pour choisir le dossier de sauvegarde
        self.choose_folder_button = ctk.CTkButton(
            filename_frame,
            text="Choisir dossier",
            command=self.choose_save_folder,
            width=120
        )
        self.choose_folder_button.pack(side="left", padx=5)

        # Label pour afficher le chemin de sauvegarde
        self.save_path_label = ctk.CTkLabel(filename_frame, text=f"Dossier: {self.word_save_path}")
        self.save_path_label.pack(side="left", padx=5)

        # Deuxième ligne : boutons d'export
        export_frame = ctk.CTkFrame(button_frame)
        export_frame.pack(fill="x", padx=5, pady=5)

        # Container centré pour les boutons
        button_container = ctk.CTkFrame(export_frame)
        button_container.pack(anchor="center", pady=5)

        # Boutons Word et PDF côte à côte
        self.word_button = ctk.CTkButton(
            button_container,
            text="Exporter en Word",
            command=self.export_to_word,
            width=150
        )
        self.word_button.pack(side="left", padx=10)

        self.pdf_button = ctk.CTkButton(
            button_container,
            text="Exporter en PDF",
            command=self.export_to_pdf,
            width=150
        )
        self.pdf_button.pack(side="left", padx=10)
        
        # Barre d'outils de mise en forme
        toolbar_frame = ctk.CTkFrame(result_frame)
        toolbar_frame.pack(fill="x", padx=5, pady=5)
        
        # Boutons de mise en forme
        self.bold_button = ctk.CTkButton(
            toolbar_frame,
            text="Gras",
            width=60,
            command=lambda: self.toggle_text_style("bold")
        )
        self.bold_button.pack(side="left", padx=2)
        
        self.italic_button = ctk.CTkButton(
            toolbar_frame,
            text="Italique",
            width=60,
            command=lambda: self.toggle_text_style("italic")
        )
        self.italic_button.pack(side="left", padx=2)
        
        # Boutons d'alignement
        align_frame = ctk.CTkFrame(toolbar_frame)
        align_frame.pack(side="left", padx=10)
        
        for align, text in [("center", "↔"), ("right", "→"), ("justify", "≡")]:
            btn = ctk.CTkButton(
                align_frame,
                text=text,
                width=30,
                command=lambda a=align: self.set_text_alignment(a)
            )
            btn.pack(side="left", padx=1)
        
        # Menu déroulant pour l'interligne
        spacing_label = ctk.CTkLabel(toolbar_frame, text="Interligne:")
        spacing_label.pack(side="left", padx=10)
        
        self.spacing_var = ctk.StringVar(value="1.0")
        spacing_menu = ctk.CTkOptionMenu(
            toolbar_frame,
            values=["1.0", "1.15", "1.5", "2.0"],
            variable=self.spacing_var,
            command=self.set_line_spacing,
            width=70
        )
        spacing_menu.pack(side="left", padx=5)
        
        # Zone de texte du résultat avec apparence Word
        self.result_text = ctk.CTkTextbox(
            result_frame,
            height=400,
            wrap="word"
        )
        self.result_text.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Configurer l'apparence Word
        self.result_text._textbox.configure(
            relief="solid",
            borderwidth=1,
            padx=40,  # Marges gauche/droite
            pady=40,   # Marges haut/bas
            font=("Times New Roman", 12)
        )
    
    def create_input_section(self, parent):
        # Section des informations
        info_frame = ctk.CTkFrame(parent)
        info_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(
            info_frame,
            text="Informations",
            font=("Arial", 16, "bold")
        ).pack(anchor="w", pady=10)
        
        # Créer les champs d'entrée
        self.company_entry = self.create_input_field(info_frame, "Entreprise :")
        self.position_entry = self.create_input_field(info_frame, "Poste :")
        self.duration_entry = self.create_input_field(info_frame, "Durée :")
        self.start_date_entry = self.create_input_field(info_frame, "Date de début :")
        
        # Ajouter le champ pour la date du jour
        self.today_date_entry = self.create_input_field(info_frame, "Date du jour :")
        
        # Zone pour le paragraphe personnalisé
        self.create_custom_text_frame(info_frame)
        
        # Bouton pour valider les informations
        self.validate_button = ctk.CTkButton(
            info_frame,
            text="Valider les informations",
            command=self.validate_and_update_info,
            font=("Arial", 12, "bold")
        )
        self.validate_button.pack(pady=10)
        
        # Frame pour le nom du fichier et le bouton d'export
        export_frame = ctk.CTkFrame(info_frame)
        export_frame.pack(fill="x", pady=(10, 0))

        # Label et champ pour le nom du fichier
        filename_label = ctk.CTkLabel(export_frame, text="Nom du fichier:")
        filename_label.pack(side="left", padx=5)
        
        self.filename_entry = ctk.CTkEntry(export_frame, width=200)
        self.filename_entry.pack(side="left", padx=5)

        # Bouton pour choisir le dossier de sauvegarde
        self.choose_folder_button = ctk.CTkButton(
            export_frame,
            text="Choisir dossier",
            command=self.choose_save_folder
        )
        self.choose_folder_button.pack(side="left", padx=5)

        # Label pour afficher le chemin de sauvegarde
        self.save_path_label = ctk.CTkLabel(export_frame, text=f"Dossier: {self.word_save_path}")
        self.save_path_label.pack(side="left", padx=5)
    
    def create_custom_text_frame(self, parent):
        """Créer le cadre pour le texte personnalisé."""
        custom_frame = ctk.CTkFrame(parent)
        custom_frame.pack(fill="x", padx=10, pady=5)
        
        # Label
        custom_label = ctk.CTkLabel(custom_frame, text="Paragraphe personnalisé :")
        custom_label.pack(side="top", anchor="w")
        
        # Frame pour les contrôles
        controls_frame = ctk.CTkFrame(custom_frame)
        controls_frame.pack(fill="x", pady=(5, 0))
        
        # Menu déroulant pour les messages prérédigés
        self.custom_templates = {
            "Sélectionner un message...": "",
            "Enthousiasme": "Je suis particulièrement enthousiaste à l'idée de rejoindre votre équipe et de contribuer activement à vos projets innovants.",
            "Disponibilité": "Je suis disponible immédiatement et prêt(e) à m'investir pleinement dans ce nouveau défi professionnel.",
            "Motivation": "Votre entreprise correspond parfaitement à mes aspirations professionnelles et je suis convaincu(e) de pouvoir apporter une réelle valeur ajoutée à votre équipe.",
            "Expertise": "Mon expertise dans ce domaine, acquise au fil de mes expériences, sera un atout précieux pour ce poste.",
            "Adaptation": "Ma capacité d'adaptation et mon envie d'apprendre me permettront de m'intégrer rapidement au sein de votre équipe.",
        }
        
        # Frame pour le menu et les boutons de gestion
        template_controls = ctk.CTkFrame(controls_frame)
        template_controls.pack(side="left", fill="x", expand=True)
        
        self.custom_template_var = tkinter.StringVar(value="Sélectionner un message...")
        self.custom_template_menu = ctk.CTkOptionMenu(
            template_controls,
            values=list(self.custom_templates.keys()),
            variable=self.custom_template_var,
            command=self.update_custom_text,
            width=200
        )
        self.custom_template_menu.pack(side="left", padx=5)
        
        # Boutons de gestion des modèles
        edit_button = ctk.CTkButton(
            template_controls,
            text="✎",
            width=30,
            command=self.edit_template
        )
        edit_button.pack(side="left", padx=2)
        
        rename_button = ctk.CTkButton(
            template_controls,
            text="✍",
            width=30,
            command=self.rename_template
        )
        rename_button.pack(side="left", padx=2)
        
        delete_button = ctk.CTkButton(
            template_controls,
            text="✕",
            width=30,
            command=self.delete_template,
            fg_color="darkred",
            hover_color="red"
        )
        delete_button.pack(side="left", padx=2)
        
        add_button = ctk.CTkButton(
            template_controls,
            text="+",
            width=30,
            command=self.add_template
        )
        add_button.pack(side="left", padx=2)
        
        # Bouton pour effacer
        clear_button = ctk.CTkButton(
            controls_frame,
            text="Effacer",
            width=70,
            command=lambda: self.custom_text.delete("1.0", "end")
        )
        clear_button.pack(side="right", padx=5)
        
        # Zone de texte personnalisé
        self.custom_text = ctk.CTkTextbox(custom_frame, height=100)
        self.custom_text.pack(fill="x", pady=5)
    
    def add_template(self):
        """Ajouter un nouveau modèle."""
        try:
            # Demander le nom du nouveau modèle
            dialog = ctk.CTkInputDialog(
                text="Nom du nouveau modèle :",
                title="Nouveau modèle"
            )
            name = dialog.get_input()
            
            if name:
                if name in self.custom_templates:
                    self.show_status("Ce nom de modèle existe déjà.", is_error=True)
                    return
                    
                # Ajouter le modèle avec le texte actuel
                current_text = self.custom_text.get("1.0", "end-1c")
                self.custom_templates[name] = current_text
                
                # Mettre à jour le menu déroulant
                menu_items = list(self.custom_templates.keys())
                menu_items.sort()  # Trier les items par ordre alphabétique
                if "Sélectionner un message..." in menu_items:
                    menu_items.remove("Sélectionner un message...")
                    menu_items.insert(0, "Sélectionner un message...")
                
                self.custom_template_menu.configure(values=menu_items)
                self.custom_template_var.set(name)
                
                # Sauvegarder immédiatement
                self.save_custom_templates()
                self.show_status(f"Nouveau modèle ajouté : {name}")
                
        except Exception as e:
            self.show_status(f"Erreur lors de l'ajout du modèle : {str(e)}", is_error=True)
    
    def edit_template(self):
        """Sauvegarder le texte actuel comme nouveau modèle."""
        try:
            # Récupérer le modèle sélectionné
            name = self.custom_template_var.get()
            if name == "Sélectionner un message...":
                self.show_status("Veuillez sélectionner un modèle à modifier.", is_error=True)
                return
            
            # Récupérer le texte actuel
            current_text = self.custom_text.get("1.0", "end-1c")
            
            # Mettre à jour le modèle avec le texte actuel
            self.custom_templates[name] = current_text
            
            # Sauvegarder les modèles
            self.save_custom_templates()
            
            # Afficher une confirmation
            self.show_status("Le modèle a été mis à jour avec succès.")
            
        except Exception as e:
            self.show_status(f"Erreur lors de la modification du modèle : {str(e)}", is_error=True)
    
    def rename_template(self):
        """Renommer le modèle sélectionné."""
        try:
            # Récupérer le modèle sélectionné
            old_name = self.custom_template_var.get()
            if old_name == "Sélectionner un message...":
                self.show_status("Veuillez sélectionner un modèle à renommer.", is_error=True)
                return
            
            # Demander le nouveau nom
            dialog = ctk.CTkInputDialog(
                text="Nouveau nom du modèle :",
                title="Renommer le modèle"
            )
            new_name = dialog.get_input()
            
            if new_name:
                if new_name in self.custom_templates:
                    self.show_status("Ce nom de modèle existe déjà.", is_error=True)
                    return
                    
                # Copier le contenu sous le nouveau nom
                self.custom_templates[new_name] = self.custom_templates[old_name]
                # Supprimer l'ancien
                del self.custom_templates[old_name]
                
                # Mettre à jour le menu déroulant
                menu_items = list(self.custom_templates.keys())
                menu_items.sort()  # Trier les items par ordre alphabétique
                if "Sélectionner un message..." in menu_items:
                    menu_items.remove("Sélectionner un message...")
                    menu_items.insert(0, "Sélectionner un message...")
                
                self.custom_template_menu.configure(values=menu_items)
                self.custom_template_var.set(new_name)
                
                # Sauvegarder les modèles
                self.save_custom_templates()
                
        except Exception as e:
            self.show_status(f"Erreur lors du renommage du modèle : {str(e)}", is_error=True)
    
    def delete_template(self):
        """Supprimer le modèle sélectionné."""
        try:
            # Récupérer le modèle sélectionné
            name = self.custom_template_var.get()
            if name == "Sélectionner un message...":
                self.show_status("Veuillez sélectionner un modèle à supprimer.", is_error=True)
                return
            
            # Demander confirmation
            if messagebox.askyesno("Confirmation", f"Voulez-vous vraiment supprimer le modèle '{name}' ?"):
                # Supprimer le modèle
                del self.custom_templates[name]
                
                # Mettre à jour le menu déroulant
                menu_items = list(self.custom_templates.keys())
                menu_items.sort()  # Trier les items par ordre alphabétique
                if "Sélectionner un message..." in menu_items:
                    menu_items.remove("Sélectionner un message...")
                    menu_items.insert(0, "Sélectionner un message...")
                
                self.custom_template_menu.configure(values=menu_items)
                self.custom_template_var.set("Sélectionner un message...")
                
                # Effacer le texte si c'était le modèle sélectionné
                self.custom_text.delete("1.0", "end")
                
                # Sauvegarder les modèles
                self.save_custom_templates()
                
        except Exception as e:
            self.show_status(f"Erreur lors de la suppression du modèle : {str(e)}", is_error=True)
    
    def update_custom_text(self, selection):
        """Mettre à jour le texte personnalisé avec le modèle sélectionné."""
        try:
            if selection in self.custom_templates:
                self.custom_text.delete("1.0", "end")
                self.custom_text.insert("1.0", self.custom_templates[selection])
                self.show_status(f"Modèle chargé : {selection}")
        except Exception as e:
            self.show_status(f"Erreur lors du chargement du modèle : {str(e)}", is_error=True)
    
    def create_template_section(self, parent):
        template_frame = ctk.CTkFrame(parent)
        template_frame.pack(fill="x", pady=(0, 20))
        
        # Titre de la section
        title_label = ctk.CTkLabel(
            template_frame, 
            text="2. Modèle de lettre", 
            font=("Arial", 20, "bold")
        )
        title_label.pack(pady=10)
        
        # Frame pour les boutons
        button_frame = ctk.CTkFrame(template_frame)
        button_frame.pack(fill="x", padx=20, pady=5)
        
        # Bouton pour effacer le modèle
        clear_model_button = ctk.CTkButton(
            button_frame,
            text="Effacer le modèle",
            command=self.clear_template,
            width=120
        )
        clear_model_button.pack(side="left", padx=5)
        
        # Frame pour la légende des marqueurs
        legend_frame = ctk.CTkFrame(template_frame)
        legend_frame.pack(fill="x", padx=20, pady=5)
        
        legend_label = ctk.CTkLabel(
            legend_frame,
            text="Légende des marqueurs :",
            font=("Arial", 14, "bold")
        )
        legend_label.pack(anchor="w", pady=5)
        
        # Créer un exemple pour chaque marqueur avec des descriptions très précises
        markers = [
            ("company", " Entreprise (texte surligné en VERT sera remplacé par la valeur saisie dans 'Entreprise')", "#34C759"),  # vert clair
            ("position", " Poste (texte surligné en BLEU sera remplacé par la valeur saisie dans 'Poste')", "#3498DB"),  # bleu ciel
            ("duration", " Durée (texte surligné en VIOLET sera remplacé par la valeur saisie dans 'Durée')", "#9B59B6"),  # violet clair
            ("start_date", " Date de début (texte surligné en JAUNE sera remplacé par la valeur saisie dans 'Date de début')", "#F7DC6F"),  # jaune clair
            ("today_date", " Date du jour (texte surligné en ROSE sera remplacé par la valeur saisie dans 'Date du jour')", "#FFC5C5"),  # rose clair
            ("custom", " Paragraphe personnalisé (tout le paragraphe surligné en ORANGE sera remplacé par le texte saisi dans 'Paragraphe personnalisé')", "#FFA07A"),  # orange clair
        ]
        
        # Créer des exemples colorés pour chaque type de marqueur
        for label, description, color in markers:
            example_frame = ctk.CTkFrame(legend_frame)
            example_frame.pack(anchor="w", padx=20, fill="x", pady=2)
            
            color_sample = tkinter.Label(
                example_frame, 
                text="   ",
                background=color
            )
            color_sample.pack(side="left", padx=5, pady=2)
            
            ctk.CTkLabel(example_frame, text=description, font=("Arial", 12)).pack(side="left", padx=5)
        
        # Barre d'outils pour le marquage et la mise en forme
        toolbar_frame = ctk.CTkFrame(template_frame)
        toolbar_frame.pack(fill="x", padx=20, pady=5)
        
        # Frame pour les outils de mise en forme
        format_toolbar = ctk.CTkFrame(toolbar_frame)
        format_toolbar.pack(side="left", padx=5)
        
        # Bouton gras
        self.template_bold_button = ctk.CTkButton(
            format_toolbar,
            text="Gras",
            width=60,
            command=lambda: self.toggle_text_style("bold", self.template_text)
        )
        self.template_bold_button.pack(side="left", padx=2)

        # Bouton italique
        self.template_italic_button = ctk.CTkButton(
            format_toolbar,
            text="Italique",
            width=60,
            command=lambda: self.toggle_text_style("italic", self.template_text)
        )
        self.template_italic_button.pack(side="left", padx=2)
        
        # Boutons d'alignement
        align_frame = ctk.CTkFrame(format_toolbar)
        align_frame.pack(side="left", padx=10)
        
        for align, text in [("center", "↔"), ("right", "→"), ("justify", "≡")]:
            btn = ctk.CTkButton(
                align_frame,
                text=text,
                width=30,
                command=lambda a=align: self.set_text_alignment(a, self.template_text)
            )
            btn.pack(side="left", padx=1)
        
        # Menu déroulant pour l'interligne
        spacing_label = ctk.CTkLabel(format_toolbar, text="Interligne:")
        spacing_label.pack(side="left", padx=10)
        
        self.template_spacing_var = ctk.StringVar(value="1.0")
        template_spacing_menu = ctk.CTkOptionMenu(
            format_toolbar,
            values=["1.0", "1.15", "1.5", "2.0"],
            variable=self.template_spacing_var,
            command=lambda s: self.change_line_spacing(s, self.template_text),
            width=70
        )
        template_spacing_menu.pack(side="left", padx=5)
        
        # Séparateur vertical
        separator = ctk.CTkFrame(toolbar_frame, width=2, height=30)
        separator.pack(side="left", padx=10, fill="y")
        
        # Frame pour les boutons de marquage
        marking_frame = ctk.CTkFrame(toolbar_frame)
        marking_frame.pack(side="left", fill="x", expand=True)
        
        # Label pour les marqueurs
        ctk.CTkLabel(marking_frame, text="Marquer la sélection comme :", font=("Arial", 12)).pack(side="left", padx=5)
        
        # Boutons pour les différents types de marquage avec leurs couleurs
        self.markers = [
            ("Entreprise", "company", "#34C759"),  # vert clair
            ("Poste", "position", "#3498DB"),  # bleu ciel
            ("Durée", "duration", "#9B59B6"),  # violet clair
            ("Date de début", "start_date", "#F7DC6F"),  # jaune clair
            ("Date du jour", "today_date", "#FFC5C5"),  # rose clair
            ("Paragraphe personnalisé", "custom", "#FFA07A"),  # orange clair
        ]
        
        # Créer les boutons de marquage dans un frame dédié avec défilement horizontal si nécessaire
        buttons_frame = ctk.CTkFrame(marking_frame)
        buttons_frame.pack(side="left", fill="x", expand=True)
        
        for label, tag, color in self.markers:
            btn = ctk.CTkButton(
                buttons_frame,
                text=label,
                command=lambda t=tag: self.mark_selection(t),
                width=120
            )
            btn.pack(side="left", padx=2)
        
        # Créer le widget Text avec support des tags
        self.template_text = tkinter.Text(
            template_frame,
            height=20,
            font=("Times New Roman", 12),
            wrap="word",
            background="#F9F9FA",
            foreground="#1D1E1E",
            padx=40,
            pady=40
        )
        self.template_text.pack(padx=20, pady=10, fill="x")
        
        # Configurer les tags avec les couleurs
        for _, tag, color in self.markers:
            self.template_text.tag_configure(tag, background=color)
        
        # Frame pour afficher les informations validées
        info_summary_frame = ctk.CTkFrame(template_frame)
        info_summary_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(
            info_summary_frame,
            text="Informations qui seront utilisées pour la génération :",
            font=("Arial", 14, "bold")
        ).pack(anchor="w", pady=(5, 10))
        
        self.summary_labels = {}
        for label, key in [
            ("Entreprise", "company"),
            ("Poste", "position"),
            ("Durée", "duration"),
            ("Date de début", "start_date"),
            ("Date du jour", "today_date"),
            ("Paragraphe personnalisé", "custom")
        ]:
            frame = ctk.CTkFrame(info_summary_frame)
            frame.pack(fill="x", padx=10, pady=2)
            ctk.CTkLabel(frame, text=f"{label} :", font=("Arial", 12, "bold")).pack(side="left", padx=5)
            self.summary_labels[key] = ctk.CTkLabel(frame, text="Non renseigné", font=("Arial", 12))
            self.summary_labels[key].pack(side="left", padx=5)
        
        # Bouton pour générer
        self.generate_button = ctk.CTkButton(
            template_frame,
            text="Générer la lettre",
            command=self.generate_letter,
            font=("Arial", 14, "bold"),
            height=40
        )
        self.generate_button.pack(pady=20)
    
    def create_input_field(self, parent, label):
        frame = ctk.CTkFrame(parent)
        frame.pack(fill="x", padx=10, pady=2)
        ctk.CTkLabel(frame, text=label, font=("Arial", 12, "bold")).pack(side="left", padx=5)
        entry = ctk.CTkEntry(frame, width=200, font=("Arial", 12))
        entry.pack(side="left", padx=(0, 5))
        ctk.CTkButton(
            frame,
            text="",
            width=30,
            height=30,
            command=lambda: entry.delete(0, "end"),
            fg_color="#8B0A1A",
            hover_color="#FF3737"
        ).pack(side="left")
        return entry
    
    def bind_keyboard_shortcuts(self, widget):
        # Activer les raccourcis clavier standards
        widget.bind("<Control-v>", lambda e: self.paste_text(widget))
        widget.bind("<Command-v>", lambda e: self.paste_text(widget))  # Pour Mac
        widget.bind("<Control-c>", lambda e: self.copy_text(widget))
        widget.bind("<Command-c>", lambda e: self.copy_text(widget))  # Pour Mac
    
    def copy_text(self, widget):
        try:
            selected_text = widget.get("sel.first", "sel.last")
            self.window.clipboard_clear()
            self.window.clipboard_append(selected_text)
        except:
            pass
    
    def paste_text(self, widget):
        try:
            text = self.window.clipboard_get()
            widget.insert("insert", text)
        except:
            pass
    
    def copy_result(self):
        try:
            text = self.result_text.get("1.0", "end-1c")
            self.window.clipboard_clear()
            self.window.clipboard_append(text)
            self.show_status("La lettre a été copiée dans le presse-papiers")
        except Exception as e:
            self.show_status(f"Erreur lors de la copie : {str(e)}", is_error=True)
    
    def validate_and_update_info(self):
        """Valider les informations et mettre à jour l'interface."""
        try:
            # Vérifier que tous les champs sont remplis
            fields = {
                "Entreprise": self.company_entry.get(),
                "Poste": self.position_entry.get(),
                "Durée": self.duration_entry.get(),
                "Date de début": self.start_date_entry.get()
            }
            
            empty_fields = [name for name, value in fields.items() if not value.strip()]
            if empty_fields:
                self.show_status(f"Veuillez remplir les champs suivants :\n- {chr(10).join(empty_fields)}", is_error=True)
                return False
            
            # Mettre à jour le nom du fichier Word
            company_name = self.company_entry.get().strip()
            filename = f"stage_adrien_LANGE_lettre de motivation_{company_name}.docx"
            self.filename_entry.configure(state="normal")
            self.filename_entry.delete(0, "end")
            self.filename_entry.insert(0, filename)
            self.filename_entry.configure(state="readonly")
            
            # Mettre à jour les labels de résumé
            self.summary_labels["company"].configure(text=self.company_entry.get())
            self.summary_labels["position"].configure(text=self.position_entry.get())
            self.summary_labels["duration"].configure(text=self.duration_entry.get() if self.duration_entry.get() else "Non renseigné")
            self.summary_labels["start_date"].configure(text=self.start_date_entry.get() if self.start_date_entry.get() else "Non renseigné")
            self.summary_labels["today_date"].configure(text=self.today_date_entry.get() if self.today_date_entry.get() else "Non renseigné")
            self.summary_labels["custom"].configure(text="Paragraphe saisi" if self.custom_text.get("1.0", "end-1c") else "Non renseigné")
            
            # Afficher un message de confirmation
            self.show_status("Informations enregistrées avec succès !")
            return True
            
        except Exception as e:
            self.show_status(f"Erreur lors de la validation : {str(e)}", is_error=True)
            return False
    
    def load_model(self):
        """Charger le modèle GPT4All."""
        if self.llm is None:
            try:
                from gpt4all import GPT4All
                self.llm = GPT4All("mistral-7b-instruct-v0.1.Q4_0.gguf")
                return True
            except Exception as e:
                self.show_status(f"Impossible de charger le modèle : {str(e)}", is_error=True)
                return False
        return True
    
    def process_long_text(self, text, max_length=1000):
        """Traite un texte long en le divisant en sections"""
        sections = []
        words = text.split()
        current_section = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 <= max_length:
                current_section.append(word)
                current_length += len(word) + 1
            else:
                sections.append(" ".join(current_section))
                current_section = [word]
                current_length = len(word)
        
        if current_section:
            sections.append(" ".join(current_section))
        
        return sections
    
    def mark_selection(self, tag):
        """Marquer la sélection avec un tag spécifique."""
        try:
            # Récupérer la sélection actuelle
            try:
                start = self.template_text.index("sel.first")
                end = self.template_text.index("sel.last")
            except tkinter.TclError:  # Aucune sélection
                self.show_status("Veuillez d'abord sélectionner du texte", is_error=True)
                return
            
            # Supprimer les autres marqueurs de la sélection
            for marker_tag in self.marker_colors.keys():
                self.template_text.tag_remove(marker_tag, start, end)
            
            # Appliquer le nouveau marqueur avec sa couleur
            self.template_text.tag_add(tag, start, end)
            self.template_text.tag_configure(tag, background=self.marker_colors[tag])
            
        except Exception as e:
            self.show_status(f"Erreur lors du marquage : {str(e)}", is_error=True)

    def clear_template(self):
        # Supprimer tout le texte
        self.template_text.delete("1.0", "end")
        
        # Supprimer tous les tags
        for _, tag, _ in self.markers:
            self.template_text.tag_remove(tag, "1.0", "end")
    
    def generate_letter(self):
        """Générer la lettre finale en remplaçant les parties marquées."""
        try:
            # Récupérer le texte du modèle
            template = self.template_text.get("1.0", "end-1c")
            
            # Récupérer tous les styles du modèle
            template_styles = self.get_text_styles(self.template_text)
            
            # Effacer le texte résultat précédent
            self.result_text._textbox.delete("1.0", "end")
            
            # Insérer le nouveau texte
            self.result_text._textbox.insert("1.0", template)
            
            # Appliquer l'interligne du modèle à tout le texte d'abord
            spacing_value = self.template_spacing_var.get()
            tag_name = f"spacing_{spacing_value.replace('.', '_')}"
            self.result_text._textbox.tag_configure(
                tag_name,
                spacing1=int(float(spacing_value) * 2),
                spacing3=int(float(spacing_value) * 2)
            )
            self.result_text._textbox.tag_add(tag_name, "1.0", "end")
            
            # Appliquer tous les styles du modèle
            for style in template_styles:
                # Appliquer le style au même endroit dans le résultat
                self.result_text._textbox.tag_add(style["tag"], style["start"], style["end"])
                
                # Configurer le style selon son type
                if style["tag"] == "bold":
                    self.result_text._textbox.tag_configure(style["tag"], font=("Times New Roman", 12, "bold"))
                elif style["tag"] == "italic":
                    self.result_text._textbox.tag_configure(style["tag"], font=("Times New Roman", 12, "italic"))
                elif style["tag"].startswith("align_"):
                    align = style["tag"].replace("align_", "")
                    if align == "center":
                        self.result_text._textbox.tag_configure(style["tag"], justify="center")
                    elif align == "right":
                        self.result_text._textbox.tag_configure(style["tag"], justify="right")
                    elif align == "justify":
                        self.result_text._textbox.tag_configure(style["tag"], justify="left")
                elif style["tag"].startswith("spacing_"):
                    # Récupérer la valeur d'interligne du tag (par ex: spacing_1_15 -> 1.15)
                    spacing_value = float(style["tag"].replace("spacing_", "").replace("_", "."))
                    self.result_text._textbox.tag_configure(
                        style["tag"],
                        spacing1=int(spacing_value * 2),
                        spacing3=int(spacing_value * 2)
                    )
            
            # Récupérer les plages marquées
            marked_ranges = self.get_marked_ranges()
            
            # Faire les remplacements en commençant par la fin pour ne pas perturber les positions
            for tag, start, end in reversed(marked_ranges):
                # Récupérer le texte à remplacer
                if tag == "company":
                    # Vérifier si ce remplacement est dans le titre
                    start_line = self.result_text._textbox.index(f"{start} linestart")
                    line_content = self.result_text._textbox.get(start_line, f"{start_line} lineend")
                    replacement = self.company_entry.get()
                    # Si "LETTRE DE MOTIVATION" est dans la ligne, mettre en majuscules
                    if "LETTRE DE MOTIVATION" in line_content.upper():
                        replacement = replacement.upper()
                elif tag == "position":
                    replacement = self.position_entry.get()
                elif tag == "duration":
                    replacement = self.duration_entry.get()
                elif tag == "start_date":
                    replacement = self.start_date_entry.get()
                elif tag == "today_date":
                    replacement = self.today_date_entry.get()
                elif tag == "custom":
                    replacement = self.custom_text.get("1.0", "end-1c").strip()
                else:
                    continue
                
                # Récupérer les styles à cette position
                styles_at_pos = []
                for style in template_styles:
                    if self.is_position_in_range(start, style["start"], style["end"]):
                        styles_at_pos.append(style["tag"])
                
                # Remplacer le texte
                self.result_text._textbox.delete(str(start), str(end))
                self.result_text._textbox.insert(str(start), replacement)
                
                # Réappliquer les styles sur le nouveau texte
                new_end = f"{start}+{len(replacement)}c"
                for style_tag in styles_at_pos:
                    self.result_text._textbox.tag_add(style_tag, str(start), new_end)
            
            # Afficher un message de succès
            self.show_status("La lettre a été générée avec succès !")
            
        except Exception as e:
            self.show_status(f"Impossible de générer la lettre : {str(e)}", is_error=True)

    def is_position_in_range(self, pos, start, end):
        """Vérifie si une position est dans une plage donnée."""
        try:
            return (self.result_text._textbox.compare(pos, ">=", start) and 
                    self.result_text._textbox.compare(pos, "<=", end))
        except:
            return False

    def _apply_appearance_mode(self, color):
        """Apply the current appearance mode to a color."""
        if isinstance(color, tuple):
            return color[self._get_appearance_mode()]
        else:
            return color

    def _get_appearance_mode(self):
        """Get the current appearance mode (0 for light, 1 for dark)."""
        return 0 if ctk.get_appearance_mode() == "Light" else 1

    def save_last_session(self):
        """Sauvegarder les données de la session actuelle."""
        try:
            # Récupérer les données des champs
            data = {
                "company": self.company_entry.get(),
                "position": self.position_entry.get(),
                "duration": self.duration_entry.get(),
                "start_date": self.start_date_entry.get(),
                "today_date": self.today_date_entry.get(),
                "template": self.template_text.get("1.0", "end-1c"),
                "text_styles": self.get_text_styles(self.template_text),
                "custom": self.custom_text.get("1.0", "end-1c"),
                "template_style": self.template_style,
                "custom_templates": self.custom_templates,  # Sauvegarde des modèles personnalisés
                "word_save_path": self.word_save_path
            }
            
            # Sauvegarder dans un fichier JSON
            with open(self.save_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
                
        except Exception as e:
            self.show_status(f"Erreur lors de la sauvegarde : {str(e)}", is_error=True)

    def get_text_styles(self, widget):
        """Récupérer tous les styles appliqués au texte."""
        styles = []
        
        # Liste de tous les tags à sauvegarder
        all_tags = [
            "bold", "italic",  # Styles de texte
            "align_center", "align_right", "align_justify",  # Alignements
            "company", "position", "duration", "start_date", "today_date", "custom"  # Marqueurs
        ]
        
        # Ajouter les tags d'interligne
        spacing_tags = [tag for tag in widget.tag_names() if tag.startswith("spacing_")]
        all_tags.extend(spacing_tags)
        
        # Pour chaque tag
        for tag in all_tags:
            # Récupérer toutes les plages où ce tag est appliqué
            ranges = widget.tag_ranges(tag)
            
            # Convertir les plages en paires début-fin
            for i in range(0, len(ranges), 2):
                start = ranges[i]
                end = ranges[i + 1]
                
                styles.append({
                    "tag": tag,
                    "start": str(start),
                    "end": str(end)
                })
        
        return styles

    def apply_saved_styles(self, widget, styles):
        """Appliquer les styles sauvegardés au texte."""
        try:
            # Couleurs des marqueurs
            marker_colors = {
                "company": "#34C759",  # vert clair
                "position": "#3498DB",  # bleu ciel
                "duration": "#9B59B6",  # violet clair
                "start_date": "#F7DC6F",  # jaune clair
                "today_date": "#FFC5C5",  # rose clair
                "custom": "#FFA07A"  # orange clair
            }
            
            for style in styles:
                tag = style["tag"]
                start = style["start"]
                end = style["end"]
                
                # Appliquer le tag
                widget.tag_add(tag, start, end)
                
                # Configurer le style selon le type de tag
                if tag == "bold":
                    widget.tag_configure(tag, font=("Times New Roman", 12, "bold"))
                elif tag == "italic":
                    widget.tag_configure(tag, font=("Times New Roman", 12, "italic"))
                elif tag == "align_center":
                    widget.tag_configure(tag, justify="center")
                elif tag == "align_right":
                    widget.tag_configure(tag, justify="right")
                elif tag == "align_justify":
                    widget.tag_configure(tag, justify="left")
                elif tag in marker_colors:  # Pour les marqueurs
                    widget.tag_configure(tag, background=marker_colors[tag])
                elif tag.startswith("spacing_"):
                    # Récupérer la valeur d'interligne du tag (par ex: spacing_1_15 -> 1.15)
                    spacing_value = float(tag.replace("spacing_", "").replace("_", "."))
                    widget.tag_configure(
                        tag,
                        spacing1=int(spacing_value * 2),
                        spacing3=int(spacing_value * 2)
                    )
                    # Mettre à jour la valeur dans le menu déroulant
                    if isinstance(widget, tkinter.Text):
                        # Si c'est le widget de modèle
                        if widget == self.template_text._textbox:
                            self.template_spacing_var.set(str(spacing_value))
                        # Si c'est le widget de résultat
                        elif widget == self.result_text._textbox:
                            self.spacing_var.set(str(spacing_value))
        
        except Exception as e:
            self.show_status(f"Erreur lors de l'application des styles : {str(e)}", is_error=True)

    def load_last_session(self):
        """Charger les données de la dernière session."""
        try:
            if os.path.exists(self.save_file):
                with open(self.save_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Restaurer le style
                if "template_style" in data:
                    self.template_style = data["template_style"]
                    self.apply_template_style()
                
                # Restaurer les données
                self.company_entry.insert(0, data.get("company", ""))
                self.position_entry.insert(0, data.get("position", ""))
                self.duration_entry.insert(0, data.get("duration", ""))
                self.start_date_entry.insert(0, data.get("start_date", ""))
                self.today_date_entry.insert(0, data.get("today_date", ""))
                self.template_text.insert("1.0", data.get("template", ""))
                self.custom_text.insert("1.0", data.get("custom", ""))
                
                # Restaurer les styles
                self.apply_saved_styles(self.template_text, data.get("text_styles", []))
                
                # Restaurer les modèles personnalisés
                if "custom_templates" in data:
                    self.custom_templates = data["custom_templates"]
                    # Mettre à jour le menu déroulant avec les modèles chargés
                    menu_items = list(self.custom_templates.keys())
                    menu_items.sort()  # Trier les items par ordre alphabétique
                    if "Sélectionner un message..." in menu_items:
                        menu_items.remove("Sélectionner un message...")
                        menu_items.insert(0, "Sélectionner un message...")
                    
                    self.custom_template_menu.configure(values=menu_items)
                    self.custom_template_var.set("Sélectionner un message...")
                
                # Restaurer le chemin de sauvegarde Word
                if "word_save_path" in data:
                    self.word_save_path = data["word_save_path"]
                    if hasattr(self, 'save_path_label'):
                        self.save_path_label.configure(text=f"Dossier: {self.word_save_path}")
                
        except Exception as e:
            self.show_status(f"Erreur lors du chargement : {str(e)}", is_error=True)

    def setup_auto_save(self):
        """Configure la sauvegarde automatique pour tous les champs."""
        # Pour les champs de texte simple
        self.company_entry.bind('<KeyRelease>', lambda e: self.save_last_session())
        self.position_entry.bind('<KeyRelease>', lambda e: self.save_last_session())
        self.duration_entry.bind('<KeyRelease>', lambda e: self.save_last_session())
        self.start_date_entry.bind('<KeyRelease>', lambda e: self.save_last_session())
        self.today_date_entry.bind('<KeyRelease>', lambda e: self.save_last_session())
        
        # Pour les zones de texte
        self.custom_text.bind('<KeyRelease>', lambda e: self.save_last_session())
        self.template_text.bind('<KeyRelease>', lambda e: self.save_last_session())
        
        # Sauvegarder aussi quand on modifie les marqueurs
        self.template_text.bind('<<Selection>>', lambda e: self.save_last_session())
    
    def on_closing(self):
        """Appelé quand l'application se ferme."""
        self.save_last_session()
        self.window.destroy()
    
    def run(self):
        self.window.mainloop()

    def export_to_word(self):
        """Exporter la lettre générée en document Word."""
        try:
            if not self.validate_and_update_info():
                return

            # Utiliser le nom de fichier généré
            filename = self.filename_entry.get()
            if not filename:
                self.show_status("Le nom du fichier n'est pas valide.", is_error=True)
                return

            # Construire le chemin complet du fichier
            if not filename.lower().endswith('.docx'):
                filename += '.docx'
            file_path = os.path.join(self.word_save_path, filename)

            # Demander confirmation si le fichier existe déjà
            if os.path.exists(file_path):
                if not messagebox.askyesno("Confirmation", "Le fichier existe déjà. Voulez-vous le remplacer?"):
                    return

            # Créer le document Word
            doc = Document()
            
            # Configuration des marges (en millimètres)
            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(25)      # 2.5 cm en haut
                section.bottom_margin = Mm(25)   # 2.5 cm en bas
                section.left_margin = Mm(20)     # 2 cm à gauche
                section.right_margin = Mm(20)    # 2 cm à droite
            
            # Récupérer le texte et les styles
            text = self.result_text._textbox.get("1.0", "end-1c")
            styles = self.get_text_styles(self.result_text._textbox)
            
            # Traiter le texte ligne par ligne
            lines = text.split('\n')
            current_pos = "1.0"
            
            # Pour chaque ligne
            for i, line in enumerate(lines):
                # Créer un nouveau paragraphe pour la ligne
                p = doc.add_paragraph()
                
                # Position de fin de la ligne actuelle
                try:
                    line_end = self.result_text._textbox.index(f"{current_pos} lineend")
                except:
                    current_pos = f"{int(current_pos.split('.')[0])+1}.0"
                    continue
                
                # Vérifier l'alignement et l'interligne du paragraphe
                for style in styles:
                    if style["tag"].startswith("align_"):
                        start_idx = style["start"]
                        end_idx = style["end"]
                        if self.is_position_in_range(current_pos, start_idx, end_idx):
                            align = style["tag"].replace("align_", "")
                            if align == "center":
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif align == "right":
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif align == "justify":
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    elif style["tag"].startswith("spacing_"):
                        start_idx = style["start"]
                        end_idx = style["end"]
                        if self.is_position_in_range(current_pos, start_idx, end_idx):
                            # Extraire la valeur d'interligne du tag (par ex: spacing_1_15 -> 1.15)
                            line_spacing = float(style["tag"].replace("spacing_", "").replace("_", "."))
                
                # Appliquer l'interligne
                p.paragraph_format.line_spacing = 1.0  # Valeur par défaut
                
                # Si la ligne est vide, ajouter un paragraphe vide avec l'espacement
                if not line.strip():
                    p.paragraph_format.space_after = Pt(12)  # Espacement après le paragraphe
                    current_pos = self.result_text._textbox.index(f"{current_pos}+1l")
                    continue
                
                # Traiter chaque caractère de la ligne
                current_char = current_pos
                while self.result_text._textbox.compare(current_char, "<", line_end):
                    try:
                        char = self.result_text._textbox.get(current_char)
                        if char == '\n':
                            break
                        
                        # Vérifier les styles à cette position
                        is_bold = False
                        is_italic = False
                        for style in styles:
                            if style["tag"] in ["bold", "italic"]:
                                start_idx = style["start"]
                                end_idx = style["end"]
                                if self.is_position_in_range(current_char, start_idx, end_idx):
                                    if style["tag"] == "bold":
                                        is_bold = True
                                    elif style["tag"] == "italic":
                                        is_italic = True
                        
                        # Ajouter le caractère avec ses styles
                        run = p.add_run(char)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11.5)
                        run.bold = is_bold
                        run.italic = is_italic
                        
                        # Passer au caractère suivant
                        current_char = self.result_text._textbox.index(f"{current_char}+1c")
                    except:
                        # En cas d'erreur sur un caractère, passer au suivant
                        current_char = self.result_text._textbox.index(f"{current_char}+1c")
                        continue
                
                # Configurer l'espacement du paragraphe
                p.paragraph_format.space_before = Pt(0)  # Pas d'espace avant
                p.paragraph_format.space_after = Pt(12)  # Espacement après le paragraphe
                p.paragraph_format.line_spacing = 1.0  # Interligne
                
                # Passer à la ligne suivante
                current_pos = self.result_text._textbox.index(f"{current_pos}+1l")
            
            # Sauvegarder le document
            doc.save(file_path)
            
            self.show_status(f"Document Word enregistré : {filename}")
            
        except Exception as e:
            self.show_status(f"Erreur lors de l'export : {str(e)}", is_error=True)

    def is_position_in_range(self, pos, start, end):
        """Vérifie si une position est dans une plage donnée."""
        try:
            return (self.result_text._textbox.compare(pos, ">=", start) and 
                    self.result_text._textbox.compare(pos, "<=", end))
        except:
            return False

    def export_to_pdf(self):
        """Exporter la lettre générée en PDF optimisé pour l'impression."""
        try:
            # Demander à l'utilisateur où sauvegarder le PDF
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                title="Enregistrer le PDF sous"
            )
            
            if not file_path:
                return
            
            # Créer un document Word temporaire
            temp_docx = "temp_letter.docx"
            doc = Document()
            
            # Configuration des marges
            section = doc.sections[0]
            section.page_height = Mm(297)  # A4
            section.page_width = Mm(210)   # A4
            section.left_margin = Mm(25)
            section.right_margin = Mm(25)
            section.top_margin = Mm(25)
            section.bottom_margin = Mm(25)
            
            def add_paragraph_with_style(text, alignment=None, bold=False, first_line_indent=None):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11.5)
                run.bold = bold
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if alignment:
                    p.alignment = alignment
                if first_line_indent:
                    p.paragraph_format.first_line_indent = first_line_indent
                return p
            
            # Informations de l'expéditeur
            for info in [
                self.nom_prenom_entry.get(),
                self.adresse_entry.get(),
                f"{self.code_postal_entry.get()} {self.ville_entry.get()}",
                self.telephone_entry.get(),
                self.email_entry.get()
            ]:
                if info.strip():
                    add_paragraph_with_style(info)
            
            # Espace
            doc.add_paragraph()
            
            # Informations du destinataire
            for info in [
                self.entreprise_entry.get(),
                self.adresse_entreprise_entry.get(),
                f"{self.code_postal_entreprise_entry.get()} {self.ville_entreprise_entry.get()}"
            ]:
                if info.strip():
                    add_paragraph_with_style(info)
            
            # Espace
            doc.add_paragraph()
            
            # Ville et date
            ville_date = self.ville_entry.get()
            date = self.today_date_entry.get()
            if ville_date and date:
                add_paragraph_with_style(
                    f"{ville_date}, le {date}",
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT
                )
            
            # Espace
            doc.add_paragraph()
            
            # Objet
            objet = self.objet_entry.get()
            if objet:
                p = doc.add_paragraph()
                run_label = p.add_run("Objet : ")
                run_label.font.name = 'Times New Roman'
                run_label.font.size = Pt(11.5)
                run_label.bold = True
                
                run_text = p.add_run(objet)
                run_text.font.name = 'Times New Roman'
                run_text.font.size = Pt(11.5)
                
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            
            # Espace
            doc.add_paragraph()
            
            # Formule de politesse début
            add_paragraph_with_style("Madame, Monsieur,")
            
            # Espace
            doc.add_paragraph()
            
            # Contenu de la lettre
            content = self.result_text.get("1.0", "end-1c")
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Mm(10)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    
                    # Traiter le texte caractère par caractère pour préserver les styles
                    current_text = ""
                    is_bold = False
                    is_italic = False
                    
                    for char in para:
                        if char == '\n':
                            if current_text:
                                run = p.add_run(current_text)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11.5)
                                run.bold = is_bold
                                run.italic = is_italic
                                current_text = ""
                        else:
                            current_text += char
                    
                    if current_text:
                        run = p.add_run(current_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11.5)
                        run.bold = is_bold
                        run.italic = is_italic
            
            # Formule de politesse fin
            add_paragraph_with_style(
                "\nJe vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées."
            )
            
            # Signature
            add_paragraph_with_style(
                f"\n{self.nom_prenom_entry.get()}",
                alignment=WD_ALIGN_PARAGRAPH.RIGHT
            )
            
            # Sauvegarder le document
            doc.save(temp_docx)
            
            # Convertir en PDF
            convert(temp_docx, file_path)
            
            self.show_status("Le PDF a été généré avec succès !")
            
        except Exception as e:
            self.show_status(f"Erreur lors de la génération du PDF : {str(e)}", is_error=True)
            
        finally:
            # Supprimer le fichier Word temporaire
            if os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                except:
                    pass
    
    def get_marked_ranges(self):
        """Récupérer les plages de texte marquées."""
        try:
            # Récupérer tous les tags
            tags = self.template_text.tag_names()
            
            # Filtre pour ne garder que les tags de marqueurs
            marker_tags = [tag for tag in tags if tag in self.marker_colors.keys()]
            
            # Récupérer les plages pour chaque tag
            ranges = []
            for tag in marker_tags:
                tag_ranges = self.template_text.tag_ranges(tag)
                if tag_ranges:
                    for i in range(0, len(tag_ranges), 2):
                        start = tag_ranges[i]
                        end = tag_ranges[i + 1]
                        ranges.append((tag, start, end))
            
            return ranges
            
        except Exception as e:
            self.show_status(f"Erreur lors de la récupération des plages marquées : {str(e)}", is_error=True)
            return []
    
    def get_marked_range(self, text_widget, tag_name):
        """Récupère la plage de texte marquée avec un tag spécifique."""
        try:
            # Récupérer tous les ranges pour ce tag
            ranges = text_widget.tag_ranges(tag_name)
            if not ranges:
                return None
            
            # Convertir les indices en positions de ligne
            ranges_list = []
            for i in range(0, len(ranges), 2):
                start = text_widget.index(ranges[i])
                end = text_widget.index(ranges[i + 1])
                ranges_list.append((start, end))
            
            return ranges_list
            
        except Exception as e:
            self.show_status(f"Erreur lors de la récupération des ranges marqués : {str(e)}", is_error=True)
            return None

    def toggle_text_style(self, style, widget=None):
        """Activer/désactiver un style de texte."""
        try:
            # Si aucun widget n'est spécifié, utiliser result_text par défaut
            if widget is None:
                widget = self.result_text._textbox
            elif isinstance(widget, ctk.CTkTextbox):
                widget = widget._textbox
            
            # Vérifier s'il y a une sélection
            try:
                start = widget.index("sel.first")
                end = widget.index("sel.last")
            except tkinter.TclError:
                self.show_status("Veuillez sélectionner du texte d'abord.", is_error=True)
                return
            
            # Si c'est un style d'alignement
            if style.startswith("align_"):
                # Supprimer tous les autres alignements
                for align in ["align_center", "align_right", "align_justify"]:
                    widget.tag_remove(align, start, end)
                    if align in self.style_buttons:
                        self.style_buttons[align].configure(fg_color=("gray86", "gray17"))
                
                # Vérifier si le style est déjà appliqué
                ranges = widget.tag_ranges(style)
                is_style_active = False
                for i in range(0, len(ranges), 2):
                    tag_start = ranges[i]
                    tag_end = ranges[i+1]
                    if (widget.compare(start, ">=", tag_start) and 
                        widget.compare(end, "<=", tag_end)):
                        is_style_active = True
                        break
                
                # Si le style est actif, le supprimer
                if is_style_active:
                    widget.tag_remove(style, start, end)
                    if style in self.style_buttons:
                        self.style_buttons[style].configure(fg_color=("gray86", "gray17"))
                # Sinon, l'appliquer
                else:
                    widget.tag_add(style, start, end)
                    if style in self.style_buttons:
                        self.style_buttons[style].configure(fg_color="#1F538D")
                
                # Configurer le style
                if style == "align_center":
                    widget.tag_configure(style, justify="center")
                elif style == "align_right":
                    widget.tag_configure(style, justify="right")
                elif style == "align_justify":
                    widget.tag_configure(style, justify="left")
            
            # Si c'est un style de texte (gras, italique)
            else:
                # Vérifier si le style est déjà appliqué sur la sélection
                ranges = widget.tag_ranges(style)
                is_style_active = False
                
                # Vérifier si toute la sélection est déjà en gras
                if ranges:
                    for i in range(0, len(ranges), 2):
                        tag_start = ranges[i]
                        tag_end = ranges[i+1]
                        if (widget.compare(start, ">=", tag_start) and 
                            widget.compare(end, "<=", tag_end)):
                            is_style_active = True
                            break
                
                # Configurer le style
                if style == "bold":
                    font = ("Times New Roman", 12, "bold")
                elif style == "italic":
                    font = ("Times New Roman", 12, "italic")
                
                # Si le style est actif, le supprimer
                if is_style_active:
                    widget.tag_remove(style, start, end)
                    if style in self.style_buttons:
                        self.style_buttons[style].configure(fg_color=("gray86", "gray17"))
                # Sinon, l'appliquer
                else:
                    widget.tag_configure(style, font=font)
                    widget.tag_add(style, start, end)
                    if style in self.style_buttons:
                        self.style_buttons[style].configure(fg_color="#1F538D")
            
        except Exception as e:
            self.show_status(f"Erreur lors du changement de style : {str(e)}", is_error=True)

    def get_active_widget(self):
        """Récupérer le widget de texte actif."""
        try:
            # Récupérer le focus
            focused = self.main_frame.focus_get()
            
            # Si c'est le modèle
            if focused == self.template_text._textbox:
                return self.template_text
            # Si c'est le résultat
            elif focused == self.result_text._textbox:
                return self.result_text
            # Par défaut, utiliser le modèle
            else:
                return self.template_text
                
        except Exception as e:
            self.show_status(f"Erreur lors de la récupération du widget actif : {str(e)}", is_error=True)
            return self.template_text

    def set_text_alignment(self, alignment, widget=None):
        """Définir l'alignement du texte"""
        try:
            # Si aucun widget n'est spécifié, utiliser result_text par défaut
            if widget is None:
                widget = self.result_text._textbox
            elif isinstance(widget, ctk.CTkTextbox):
                widget = widget._textbox
            
            # Vérifier s'il y a une sélection
            try:
                start = widget.index("sel.first")
                end = widget.index("sel.last")
            except tkinter.TclError:
                return
            
            # Supprimer les autres alignements de la sélection
            for align in ["align_center", "align_right", "align_justify"]:
                widget.tag_remove(align, start, end)
            
            # Appliquer le nouvel alignement à la sélection
            tag_name = f"align_{alignment}"
            widget.tag_add(tag_name, start, end)
            
            if alignment == "center":
                widget.tag_configure(tag_name, justify="center")
            elif alignment == "right":
                widget.tag_configure(tag_name, justify="right")
            elif alignment == "justify":
                widget.tag_configure(tag_name, justify="left")
            
        except Exception as e:
            self.show_status(f"Erreur lors du changement d'alignement : {str(e)}", is_error=True)

    def set_line_spacing(self, spacing, widget=None):
        """Définir l'interligne"""
        try:
            # Si aucun widget n'est spécifié, utiliser result_text par défaut
            if widget is None:
                widget = self.result_text._textbox
            elif isinstance(widget, ctk.CTkTextbox):
                widget = widget._textbox
            
            # Convertir en float
            spacing = float(spacing)
            
            # Appliquer l'interligne
            widget.configure(spacing3=int(spacing * 10))  # spacing3 contrôle l'interligne
            
        except Exception as e:
            self.show_status(f"Erreur lors du changement d'interligne : {str(e)}", is_error=True)

    def change_line_spacing(self, value, widget=None):
        """Changer l'interligne du texte sélectionné."""
        try:
            # Si aucun widget n'est spécifié, utiliser result_text par défaut
            if widget is None:
                widget = self.result_text._textbox
            elif isinstance(widget, ctk.CTkTextbox):
                widget = widget._textbox
            
            # Vérifier s'il y a une sélection
            try:
                start = widget.index("sel.first")
                end = widget.index("sel.last")
            except tkinter.TclError:
                self.show_status("Veuillez sélectionner du texte d'abord.", is_error=True)
                return
            
            # Convertir la valeur en float
            spacing = float(value)
            
            # Supprimer les anciens tags d'interligne
            for tag in widget.tag_names():
                if tag.startswith("spacing_"):
                    widget.tag_remove(tag, start, end)
            
            # Créer un nouveau tag pour l'interligne
            tag_name = f"spacing_{value.replace('.', '_')}"
            
            # Configurer le tag
            widget.tag_configure(
                tag_name,
                spacing1=int(spacing * 2),  # Espace avant
                spacing3=int(spacing * 2)   # Espace après
            )
            
            # Appliquer le tag du début à la fin de la ligne pour chaque ligne
            start_line = widget.index(f"{start} linestart")
            end_line = widget.index(f"{end} lineend")
            widget.tag_add(tag_name, start_line, end_line)
            
        except Exception as e:
            self.show_status(f"Erreur lors du changement d'interligne : {str(e)}", is_error=True)

    def apply_template_style(self):
        """Appliquer le style importé aux widgets de texte."""
        try:
            # Configurer la police et la taille
            font_config = (self.template_style["font_name"], int(self.template_style["font_size"]))
            
            # Appliquer aux zones de texte
            # Pour CTkTextbox, on doit configurer le widget Text interne
            for widget in [self.template_text, self.custom_text, self.result_text]:
                widget._textbox.configure(font=font_config)
                
                # Configurer l'alignement
                if self.template_style["alignment"] == "center":
                    widget._textbox.tag_configure("align", justify="center")
                elif self.template_style["alignment"] == "right":
                    widget._textbox.tag_configure("align", justify="right")
                elif self.template_style["alignment"] == "justify":
                    widget._textbox.tag_configure("align", justify="left")
                else:  # left
                    widget._textbox.tag_configure("align", justify="left")
                
                # Appliquer le tag d'alignement
                widget._textbox.tag_add("align", "1.0", "end")
                
                # Configurer l'interligne (spacing)
                widget._textbox.configure(spacing3=int(self.template_style["line_spacing"] * 3))
                widget._textbox.configure(spacing2=0)  # Espace entre les lignes
                widget._textbox.configure(spacing3=int(self.template_style["line_spacing"] * 3))
            
        except Exception as e:
            self.show_status(f"Erreur lors de l'application du style : {str(e)}", is_error=True)

    def generate_word_document(self):
        """Générer le document Word avec la lettre de motivation."""
        try:
            # Créer un nouveau document Word
            doc = Document()
            
            # Style par défaut pour tout le document
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11.5)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.line_spacing = 1.0
            
            # Configurer la mise en page
            section = doc.sections[0]
            section.page_height = Mm(297)  # A4
            section.page_width = Mm(210)   # A4
            section.left_margin = Mm(25)
            section.right_margin = Mm(25)
            section.top_margin = Mm(25)
            section.bottom_margin = Mm(25)
            
            def add_paragraph_with_style(text, alignment=None, bold=False, first_line_indent=None):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11.5)
                run.bold = bold
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if alignment:
                    p.alignment = alignment
                if first_line_indent:
                    p.paragraph_format.first_line_indent = first_line_indent
                return p
            
            # Informations de l'expéditeur
            for info in [
                self.nom_prenom_entry.get(),
                self.adresse_entry.get(),
                f"{self.code_postal_entry.get()} {self.ville_entry.get()}",
                self.telephone_entry.get(),
                self.email_entry.get()
            ]:
                if info.strip():
                    add_paragraph_with_style(info)
            
            # Espace
            doc.add_paragraph()
            
            # Informations du destinataire
            for info in [
                self.entreprise_entry.get(),
                self.adresse_entreprise_entry.get(),
                f"{self.code_postal_entreprise_entry.get()} {self.ville_entreprise_entry.get()}"
            ]:
                if info.strip():
                    add_paragraph_with_style(info)
            
            # Espace
            doc.add_paragraph()
            
            # Ville et date
            ville_date = self.ville_entry.get()
            date = self.today_date_entry.get()
            if ville_date and date:
                add_paragraph_with_style(
                    f"{ville_date}, le {date}",
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT
                )
            
            # Espace
            doc.add_paragraph()
            
            # Objet
            objet = self.objet_entry.get()
            if objet:
                p = doc.add_paragraph()
                run_label = p.add_run("Objet : ")
                run_label.font.name = 'Times New Roman'
                run_label.font.size = Pt(11.5)
                run_label.bold = True
                
                run_text = p.add_run(objet)
                run_text.font.name = 'Times New Roman'
                run_text.font.size = Pt(11.5)
                
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            
            # Espace
            doc.add_paragraph()
            
            # Formule de politesse début
            add_paragraph_with_style("Madame, Monsieur,")
            
            # Espace
            doc.add_paragraph()
            
            # Contenu de la lettre
            content = self.result_text.get("1.0", "end-1c")
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Mm(10)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    
                    # Traiter le texte caractère par caractère pour préserver les styles
                    current_text = ""
                    is_bold = False
                    is_italic = False
                    
                    for char in para:
                        if char == '\n':
                            if current_text:
                                run = p.add_run(current_text)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11.5)
                                run.bold = is_bold
                                run.italic = is_italic
                                current_text = ""
                        else:
                            current_text += char
                    
                    if current_text:
                        run = p.add_run(current_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11.5)
                        run.bold = is_bold
                        run.italic = is_italic
            
            # Formule de politesse fin
            add_paragraph_with_style(
                "\nJe vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées."
            )
            
            # Signature
            add_paragraph_with_style(
                f"\n{self.nom_prenom_entry.get()}",
                alignment=WD_ALIGN_PARAGRAPH.RIGHT
            )
            
            # Sauvegarder le document
            doc.save("lettre_motivation.docx")
            
            # Convertir en PDF
            convert("lettre_motivation.docx", "lettre_motivation.pdf")
            
            self.show_status("Les documents ont été générés avec succès !")
            
        except Exception as e:
            self.show_status(f"Erreur lors de la génération des documents : {str(e)}", is_error=True)

    def load_custom_templates(self):
        """Charger les modèles personnalisés depuis le fichier."""
        try:
            if os.path.exists(self.templates_file):
                with open(self.templates_file, 'r', encoding='utf-8') as f:
                    self.custom_templates = json.load(f)
                    self.show_status(f"Modèles chargés : {list(self.custom_templates.keys())}")
            else:
                # Modèles par défaut si le fichier n'existe pas
                self.custom_templates = {
                    "Sélectionner un message...": "",
                    "Enthousiasme": "Je suis particulièrement enthousiaste à l'idée de rejoindre votre équipe et de contribuer activement à vos projets innovants.",
                    "Disponibilité": "Je suis disponible immédiatement et prêt(e) à m'investir pleinement dans ce nouveau défi professionnel.",
                    "Motivation": "Votre entreprise correspond parfaitement à mes aspirations professionnelles et je suis convaincu(e) de pouvoir apporter une réelle valeur ajoutée à votre équipe.",
                    "Expertise": "Mon expertise dans ce domaine, acquise au fil de mes expériences, sera un atout précieux pour ce poste.",
                    "Adaptation": "Ma capacité d'adaptation et mon envie d'apprendre me permettront de m'intégrer rapidement au sein de votre équipe."
                }
                # Sauvegarder les modèles par défaut
                self.save_custom_templates()
                self.show_status("Modèles par défaut créés et sauvegardés")
        except Exception as e:
            self.show_status(f"Erreur lors du chargement des modèles : {str(e)}", is_error=True)
            self.custom_templates = {"Sélectionner un message...": ""}
    
    def save_custom_templates(self):
        """Sauvegarder les modèles personnalisés dans un fichier."""
        try:
            with open(self.templates_file, 'w', encoding='utf-8') as f:
                json.dump(self.custom_templates, f, ensure_ascii=False, indent=4)
            self.show_status(f"Modèles sauvegardés : {list(self.custom_templates.keys())}")
        except Exception as e:
            self.show_status(f"Erreur lors de la sauvegarde des modèles : {str(e)}", is_error=True)
    
    def choose_save_folder(self):
        """Permet à l'utilisateur de choisir le dossier de sauvegarde pour les fichiers Word."""
        folder_path = filedialog.askdirectory(
            title="Choisir le dossier de sauvegarde",
            initialdir=self.word_save_path
        )
        if folder_path:
            self.word_save_path = folder_path
            self.save_path_label.configure(text=f"Dossier: {self.word_save_path}")
            self.save_last_session()

    def show_status(self, message, is_error=False):
        """Affiche un message de statut dans l'interface."""
        self.status_label.configure(
            text=message,
            text_color="#dc3545" if is_error else "#28a745"  # Rouge pour les erreurs, vert pour le succès
        )
        
        # Effacer le message après 3 secondes
        self.window.after(3000, lambda: self.status_label.configure(text=""))

if __name__ == "__main__":
    app = LetterGeneratorApp()
    app.run()
