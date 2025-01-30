from flask import Flask, render_template, request, jsonify, send_file
from gpt4all import GPT4All
import os
import io
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Cm, Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import json
import re
import dotenv
from dotenv import load_dotenv

# Charger les variables d'environnement
load_dotenv()

app = Flask(__name__, static_folder='templates/web')

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max-limit

# Internationalisation
TRANSLATIONS = {
    'fr': {
        'company_required': "Le champ 'Entreprise' est requis",
        'position_required': "Le champ 'Poste' est requis",
        'invalid_date': "Format de date invalide pour '{}'. Utilisez JJ/MM/AAAA",
        'invalid_duration': "Format de durée invalide. Exemple: '6 mois', '2 ans', '3 semaines'",
        'invalid_email': "Format d'email invalide",
        'invalid_phone': "Format de téléphone invalide",
        'template_name_required': "Le nom du modèle est requis",
        'invalid_template_name': "Le nom du modèle ne doit contenir que des lettres, chiffres, espaces, tirets et underscores",
        'invalid_tag': "Le tag '{}' ne doit contenir que des lettres, chiffres, tirets et underscores",
        'template_exists': "Un modèle avec le nom '{}' existe déjà",
        'template_not_found': "Aucun modèle trouvé avec le nom '{}'",
        'template_added': "Modèle ajouté avec succès",
        'template_updated': "Modèle modifié avec succès",
        'template_deleted': "Modèle supprimé avec succès",
        'error_loading': "Erreur lors du chargement : {}",
        'error_saving': "Erreur lors de la sauvegarde : {}",
        'error_generating': "Erreur lors de la génération : {}",
        'error_exporting': "Erreur lors de l'export : {}"
    },
    'en': {
        'company_required': "The 'Company' field is required",
        'position_required': "The 'Position' field is required",
        'invalid_date': "Invalid date format for '{}'. Use DD/MM/YYYY",
        'invalid_duration': "Invalid duration format. Example: '6 months', '2 years', '3 weeks'",
        'invalid_email': "Invalid email format",
        'invalid_phone': "Invalid phone format",
        'template_name_required': "Template name is required",
        'invalid_template_name': "Template name must only contain letters, numbers, spaces, dashes and underscores",
        'invalid_tag': "Tag '{}' must only contain letters, numbers, dashes and underscores",
        'template_exists': "A template with name '{}' already exists",
        'template_not_found': "No template found with name '{}'",
        'template_added': "Template added successfully",
        'template_updated': "Template updated successfully",
        'template_deleted': "Template deleted successfully",
        'error_loading': "Error while loading: {}",
        'error_saving': "Error while saving: {}",
        'error_generating': "Error while generating: {}",
        'error_exporting': "Error while exporting: {}"
    }
}

def get_translation(key, lang='fr', *args):
    """Récupère une traduction avec formatage optionnel"""
    try:
        translation = TRANSLATIONS[lang][key]
        if args:
            return translation.format(*args)
        return translation
    except (KeyError, IndexError):
        return f"Missing translation: {key}"

class ValidationError(Exception):
    """Exception personnalisée pour les erreurs de validation"""
    pass

class LocalizedValidationError(ValidationError):
    """Exception de validation avec support des traductions"""
    def __init__(self, key, lang='fr', *args):
        self.key = key
        self.lang = lang
        self.args = args
        message = get_translation(key, lang, *args)
        super().__init__(message)

class DataValidator:
    """Classe utilitaire pour la validation des données"""
    
    @staticmethod
    def validate_required(value, field_name):
        """Valide qu'un champ requis n'est pas vide"""
        if not value or not str(value).strip():
            raise ValidationError(f"Le champ '{field_name}' est requis")
        return value.strip()
    
    @staticmethod
    def validate_date(value, field_name):
        """Valide un format de date"""
        if not value:
            return None
            
        try:
            # Accepter plusieurs formats de date
            for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%d.%m.%Y"]:
                try:
                    return datetime.strptime(value, fmt)
                except ValueError:
                    continue
            raise ValidationError(f"Format de date invalide pour '{field_name}'. Utilisez JJ/MM/AAAA")
        except Exception:
            raise ValidationError(f"Format de date invalide pour '{field_name}'. Utilisez JJ/MM/AAAA")
    
    @staticmethod
    def validate_duration(value):
        """Valide le format de la durée"""
        if not value:
            return None
            
        # Formats acceptés : "6 mois", "2 ans", "18 mois", etc.
        pattern = r"^(\d+)\s*(mois|ans?|semaines?)$"
        match = re.match(pattern, value.lower())
        if not match:
            raise ValidationError("Format de durée invalide. Exemple: '6 mois', '2 ans', '3 semaines'")
        
        number, unit = match.groups()
        return f"{number} {unit}"
    
    @staticmethod
    def validate_email(value):
        """Valide une adresse email"""
        if not value:
            return None
            
        pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
        if not re.match(pattern, value):
            raise ValidationError("Format d'email invalide")
        return value
    
    @staticmethod
    def validate_phone(value):
        """Valide un numéro de téléphone"""
        if not value:
            return None
            
        # Accepter différents formats : +33612345678, 0612345678, etc.
        pattern = r"^(?:\+33|0)\d{9}$"
        if not re.match(pattern, value.replace(" ", "")):
            raise ValidationError("Format de téléphone invalide")
        return value
    
    @staticmethod
    def validate_template_name(value):
        """Valide le nom d'un modèle"""
        if not value or not value.strip():
            raise ValidationError("Le nom du modèle est requis")
        
        # Éviter les caractères spéciaux dans le nom
        pattern = r"^[a-zA-Z0-9\s_-]+$"
        if not re.match(pattern, value):
            raise ValidationError("Le nom du modèle ne doit contenir que des lettres, chiffres, espaces, tirets et underscores")
        return value.strip()
    
    @staticmethod
    def validate_tags(tags):
        """Valide une liste de tags"""
        if not tags:
            return []
            
        valid_tags = []
        for tag in tags:
            tag = tag.strip()
            if tag:
                # Éviter les caractères spéciaux dans les tags
                pattern = r"^[a-zA-Z0-9_-]+$"
                if not re.match(pattern, tag):
                    raise ValidationError(f"Le tag '{tag}' ne doit contenir que des lettres, chiffres, tirets et underscores")
                valid_tags.append(tag)
        return valid_tags

class LocalizedDataValidator(DataValidator):
    """Version localisée du validateur de données"""
    def __init__(self, lang='fr'):
        self.lang = lang
    
    def validate_required(self, value, field_name):
        """Valide qu'un champ requis n'est pas vide"""
        if not value or not str(value).strip():
            raise LocalizedValidationError(f'{field_name.lower()}_required', self.lang)
        return value.strip()
    
    def validate_date(self, value, field_name):
        """Valide un format de date"""
        if not value:
            return None
            
        try:
            # Accepter plusieurs formats de date
            for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%d.%m.%Y"]:
                try:
                    return datetime.strptime(value, fmt)
                except ValueError:
                    continue
            raise LocalizedValidationError('invalid_date', self.lang, field_name)
        except Exception:
            raise LocalizedValidationError('invalid_date', self.lang, field_name)
    
    def validate_duration(self, value):
        """Valide le format de la durée"""
        if not value:
            return None
            
        pattern = r"^(\d+)\s*(mois|ans?|semaines?|months?|years?|weeks?)$"
        match = re.match(pattern, value.lower())
        if not match:
            raise LocalizedValidationError('invalid_duration', self.lang)
        
        number, unit = match.groups()
        # Convertir les unités en français
        unit_map = {
            'month': 'mois', 'months': 'mois',
            'year': 'an', 'years': 'ans',
            'week': 'semaine', 'weeks': 'semaines'
        }
        unit = unit_map.get(unit, unit)
        return f"{number} {unit}"
    
    def validate_email(self, value):
        """Valide une adresse email"""
        if not value:
            return None
            
        pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
        if not re.match(pattern, value):
            raise LocalizedValidationError('invalid_email', self.lang)
        return value
    
    def validate_phone(self, value):
        """Valide un numéro de téléphone"""
        if not value:
            return None
            
        pattern = r"^(?:\+33|0)\d{9}$"
        if not re.match(pattern, value.replace(" ", "")):
            raise LocalizedValidationError('invalid_phone', self.lang)
        return value
    
    def validate_template_name(self, value):
        """Valide le nom d'un modèle"""
        if not value or not value.strip():
            raise LocalizedValidationError('template_name_required', self.lang)
        
        pattern = r"^[a-zA-Z0-9\s_-]+$"
        if not re.match(pattern, value):
            raise LocalizedValidationError('invalid_template_name', self.lang)
        return value.strip()
    
    def validate_tags(self, tags):
        """Valide une liste de tags"""
        if not tags:
            return []
            
        valid_tags = []
        for tag in tags:
            tag = tag.strip()
            if tag:
                pattern = r"^[a-zA-Z0-9_-]+$"
                if not re.match(pattern, tag):
                    raise LocalizedValidationError('invalid_tag', self.lang, tag)
                valid_tags.append(tag)
        return valid_tags

def validate_letter_data(data):
    """Valide les données d'une lettre de motivation"""
    validator = LocalizedDataValidator()
    
    try:
        # Valider les champs requis
        company = validator.validate_required(data.get('company'), 'Entreprise')
        position = validator.validate_required(data.get('position'), 'Poste')
        
        # Valider les dates
        start_date = validator.validate_date(data.get('start_date'), 'Date de début')
        today_date = validator.validate_date(data.get('today_date'), 'Date du jour')
        
        # Valider la durée
        duration = validator.validate_duration(data.get('duration'))
        
        # Valider les champs optionnels
        email = validator.validate_email(data.get('email'))
        phone = validator.validate_phone(data.get('phone'))
        
        return {
            'company': company,
            'position': position,
            'start_date': start_date.strftime("%d/%m/%Y") if start_date else None,
            'today_date': today_date.strftime("%d/%m/%Y") if today_date else None,
            'duration': duration,
            'email': email,
            'phone': phone
        }
        
    except LocalizedValidationError as e:
        raise LocalizedValidationError(str(e))

def validate_template_data(data):
    """Valide les données d'un modèle"""
    validator = LocalizedDataValidator()
    
    try:
        # Valider les champs requis
        name = validator.validate_template_name(data.get('name'))
        content = validator.validate_required(data.get('content'), 'Contenu')
        
        # Valider les champs optionnels
        category = data.get('category', 'General').strip() or 'General'
        tags = validator.validate_tags(data.get('tags', []))
        
        return {
            'name': name,
            'content': content,
            'category': category,
            'tags': tags
        }
        
    except LocalizedValidationError as e:
        raise LocalizedValidationError(str(e))

class LetterGenerator:
    def __init__(self):
        # Initialiser les styles par défaut
        self.template_styles = [{
            "font_name": "Times New Roman",
            "font_size": 11,
            "bold": False,
            "alignment": 0,  # Gauche
            "line_spacing": 1.15,
            "space_before": Pt(0),
            "space_after": Pt(8),
            "first_line_indent": None
        }]
        
        # Initialiser le modèle
        self.llm = None
        
        # Créer un dossier pour sauvegarder les données si nécessaire
        self.save_dir = os.path.join(os.path.expanduser("~"), ".lettre_motivation_ai")
        os.makedirs(self.save_dir, exist_ok=True)
        self.templates_file = os.path.join(self.save_dir, "custom_templates.json")
        
        # Initialiser le dictionnaire des modèles
        self.custom_templates = {}
        
        # Charger les modèles personnalisés
        self.load_custom_templates()
        
        # Charger le modèle GPT
        self.load_model()
        
        # Ajouter les templates par défaut
        self.default_templates = {
            "Enthousiasme": "Je suis particulièrement enthousiaste à l'idée de rejoindre votre équipe et de contribuer activement à vos projets innovants.",
            "Disponibilité": "Je suis disponible immédiatement et prêt(e) à m'investir pleinement dans ce nouveau défi professionnel.",
            "Motivation": "Votre entreprise correspond parfaitement à mes aspirations professionnelles et je suis convaincu(e) de pouvoir apporter une réelle valeur ajoutée à votre équipe.",
            "Expertise": "Mon expertise dans ce domaine, acquise au fil de mes expériences, sera un atout précieux pour ce poste.",
            "Adaptation": "Ma capacité d'adaptation et mon envie d'apprendre me permettront de m'intégrer rapidement au sein de votre équipe."
        }
        self.custom_templates.update(self.default_templates)
        
        # Définir les couleurs des marqueurs
        self.marker_colors = {
            "company": "#34C759",     # Vert clair pour l'entreprise
            "position": "#3498DB",    # Bleu ciel pour le poste
            "duration": "#9B59B6",    # Violet clair pour la durée
            "start_date": "#F7DC6F",  # Jaune clair pour la date de début
            "today_date": "#FFC5C5",  # Rose clair pour la date du jour
            "custom": "#FFA07A"       # Orange clair pour le paragraphe personnalisé
        }
    
    def load_model(self):
        try:
            model_path = os.getenv("MODEL_PATH", "ggml-gpt4all-j-v1.3-groovy.bin")
            self.llm = GPT4All(model_path)
            return True
        except Exception as e:
            print(f"Erreur lors du chargement du modèle : {str(e)}")
            return False

    def load_custom_templates(self):
        try:
            if os.path.exists(self.templates_file):
                with open(self.templates_file, 'r', encoding='utf-8') as f:
                    self.custom_templates = json.load(f)
        except Exception as e:
            print(f"Erreur lors du chargement des modèles : {str(e)}")
            self.custom_templates = {}

    def generate_letter(self, data):
        if not self.llm:
            return "Erreur : Le modèle n'est pas chargé."

        # Remplacer les marqueurs par les valeurs
        template = data.get('template', '')
        for key, value in data.items():
            if key != 'template' and value:
                template = template.replace(f'[[{key}]]', value)

        # Si le template est vide, utiliser le prompt par défaut
        if not template.strip():
            prompt = f"""En tant qu'expert en rédaction de lettres de motivation, génère une lettre de motivation professionnelle et persuasive pour le poste de {data['position']} chez {data['company']}.

Informations supplémentaires :
- Type de contrat : {data['duration']}
- Date de début : {data['start_date']}
- Paragraphe personnalisé : {data['custom_paragraph']}

La lettre doit être formelle, bien structurée et mettre en avant les compétences et la motivation du candidat.
Utilise le paragraphe personnalisé pour adapter la lettre au poste et à l'entreprise.
N'inclus pas la mise en page (date, adresse, etc.) dans la réponse."""
        else:
            prompt = template

        response = self.llm.generate(
            prompt=prompt,
            max_tokens=2000,
            temp=0.7,
            top_k=40,
            top_p=0.4,
            repeat_penalty=1.18
        )

        return response

    def add_template(self, name, content):
        """Ajouter un nouveau modèle."""
        if name in self.custom_templates:
            return False, "Ce nom de modèle existe déjà."
        self.custom_templates[name] = content
        self.save_custom_templates()
        return True, "Modèle ajouté avec succès."
    
    def edit_template(self, name, content):
        """Modifier un modèle existant."""
        if name not in self.custom_templates:
            return False, "Ce modèle n'existe pas."
        self.custom_templates[name] = content
        self.save_custom_templates()
        return True, "Modèle modifié avec succès."
    
    def delete_template(self, name):
        """Supprimer un modèle."""
        if name not in self.custom_templates or name in self.default_templates:
            return False, "Impossible de supprimer ce modèle."
        del self.custom_templates[name]
        self.save_custom_templates()
        return True, "Modèle supprimé avec succès."
    
    def get_templates(self):
        """Récupérer tous les modèles."""
        return self.custom_templates

    def save_custom_templates(self):
        with open(self.templates_file, 'w', encoding='utf-8') as f:
            json.dump(self.custom_templates, f)

class Template:
    def __init__(self, name, content, category="General", tags=None):
        self.name = name
        self.content = content
        self.category = category
        self.tags = tags or []
        self.created_at = datetime.now()
        self.updated_at = datetime.now()

class LetterHistory:
    def __init__(self, company, position, content, generated_at=None):
        self.company = company
        self.position = position
        self.content = content
        self.generated_at = generated_at or datetime.now()

class TemplateManager:
    def __init__(self, save_dir):
        self.save_dir = save_dir
        self.templates_file = os.path.join(save_dir, "templates.json")
        self.history_file = os.path.join(save_dir, "history.json")
        self.templates = {}
        self.history = []
        self.load_data()

    def load_data(self):
        """Charger les modèles et l'historique"""
        try:
            if os.path.exists(self.templates_file):
                with open(self.templates_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    for template_data in data:
                        template = Template(**template_data)
                        self.templates[template.name] = template

            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.history = [LetterHistory(**item) for item in data]
        except Exception as e:
            print(f"Erreur lors du chargement des données : {e}")

    def save_data(self):
        """Sauvegarder les modèles et l'historique"""
        try:
            # Sauvegarder les modèles
            templates_data = []
            for template in self.templates.values():
                template_dict = template.__dict__
                template_dict['created_at'] = template_dict['created_at'].isoformat()
                template_dict['updated_at'] = template_dict['updated_at'].isoformat()
                templates_data.append(template_dict)

            with open(self.templates_file, 'w', encoding='utf-8') as f:
                json.dump(templates_data, f, ensure_ascii=False, indent=4)

            # Sauvegarder l'historique
            history_data = []
            for item in self.history:
                item_dict = item.__dict__
                item_dict['generated_at'] = item_dict['generated_at'].isoformat()
                history_data.append(item_dict)

            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(history_data, f, ensure_ascii=False, indent=4)

        except Exception as e:
            print(f"Erreur lors de la sauvegarde des données : {e}")

    def add_template(self, name, content, category="General", tags=None):
        """Ajouter un nouveau modèle"""
        if name in self.templates:
            raise ValueError(f"Un modèle avec le nom '{name}' existe déjà")
        
        template = Template(name, content, category, tags)
        self.templates[name] = template
        self.save_data()
        return template

    def update_template(self, name, content=None, category=None, tags=None):
        """Mettre à jour un modèle existant"""
        if name not in self.templates:
            raise ValueError(f"Aucun modèle trouvé avec le nom '{name}'")
        
        template = self.templates[name]
        if content is not None:
            template.content = content
        if category is not None:
            template.category = category
        if tags is not None:
            template.tags = tags
        template.updated_at = datetime.now()
        
        self.save_data()
        return template

    def delete_template(self, name):
        """Supprimer un modèle"""
        if name not in self.templates:
            raise ValueError(f"Aucun modèle trouvé avec le nom '{name}'")
        
        del self.templates[name]
        self.save_data()

    def get_templates_by_category(self, category):
        """Récupérer tous les modèles d'une catégorie"""
        return [t for t in self.templates.values() if t.category == category]

    def search_templates(self, query):
        """Rechercher des modèles par nom ou tags"""
        query = query.lower()
        return [t for t in self.templates.values() 
                if query in t.name.lower() 
                or any(query in tag.lower() for tag in t.tags)]

    def add_to_history(self, company, position, content):
        """Ajouter une lettre à l'historique"""
        history_item = LetterHistory(company, position, content)
        self.history.append(history_item)
        self.save_data()
        return history_item

    def get_history(self, limit=10):
        """Récupérer l'historique des lettres"""
        return sorted(self.history, 
                     key=lambda x: x.generated_at, 
                     reverse=True)[:limit]

    def search_history(self, query):
        """Rechercher dans l'historique"""
        query = query.lower()
        return [h for h in self.history 
                if query in h.company.lower() 
                or query in h.position.lower() 
                or query in h.content.lower()]

    def clear_history(self):
        """Effacer l'historique"""
        self.history = []
        self.save_data()

class DocumentExporter:
    """Classe pour gérer l'export des documents"""
    
    def __init__(self):
        self.default_font = "Times New Roman"
        self.default_font_size = 11.5
        self.default_margins = {
            "top": 25,    # mm
            "bottom": 25,
            "left": 20,
            "right": 20
        }
        self.page_size = {
            "width": 210,  # A4 en mm
            "height": 297
        }
        
    def export_to_word(self, data, file_path):
        """Exporter les données en document Word"""
        try:
            doc = Document()
            
            # Configuration des marges
            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(self.default_margins["top"])
                section.bottom_margin = Mm(self.default_margins["bottom"])
                section.left_margin = Mm(self.default_margins["left"])
                section.right_margin = Mm(self.default_margins["right"])
                
            # Informations de l'expéditeur
            self._add_sender_info(doc, data)
            
            # Espace
            doc.add_paragraph()
            
            # Informations du destinataire
            self._add_recipient_info(doc, data)
            
            # Espace
            doc.add_paragraph()
            
            # Ville et date
            if data.get("city") and data.get("date"):
                p = doc.add_paragraph()
                p.add_run(f"{data['city']}, le {data['date']}")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Espace
            doc.add_paragraph()
            
            # Objet
            if data.get("subject"):
                p = doc.add_paragraph()
                p.add_run("Objet : ").bold = True
                p.add_run(data["subject"])
            
            # Espace
            doc.add_paragraph()
            
            # Corps de la lettre
            self._add_letter_content(doc, data.get("content", ""))
            
            # Signature
            if data.get("signature"):
                p = doc.add_paragraph()
                p.add_run(data["signature"])
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Sauvegarder
            doc.save(file_path)
            return True, None
            
        except Exception as e:
            return False, str(e)
    
    def export_to_pdf(self, data, file_path):
        """Exporter les données en PDF via Word"""
        try:
            # Créer un fichier Word temporaire
            temp_docx = os.path.join(tempfile.gettempdir(), "temp_letter.docx")
            
            # Exporter en Word d'abord
            success, error = self.export_to_word(data, temp_docx)
            if not success:
                return False, error
                
            # Convertir en PDF
            convert(temp_docx, file_path)
            
            # Supprimer le fichier temporaire
            os.remove(temp_docx)
            
            return True, None
            
        except Exception as e:
            return False, str(e)
    
    def _add_sender_info(self, doc, data):
        """Ajouter les informations de l'expéditeur"""
        sender_info = [
            data.get("full_name", ""),
            data.get("address", ""),
            f"{data.get('postal_code', '')} {data.get('city', '')}".strip(),
            data.get("phone", ""),
            data.get("email", "")
        ]
        
        for info in sender_info:
            if info.strip():
                p = self._add_paragraph_with_style(doc, info)
    
    def _add_recipient_info(self, doc, data):
        """Ajouter les informations du destinataire"""
        recipient_info = [
            data.get("company", ""),
            data.get("company_address", ""),
            f"{data.get('company_postal_code', '')} {data.get('company_city', '')}".strip()
        ]
        
        for info in recipient_info:
            if info.strip():
                p = self._add_paragraph_with_style(doc, info)
    
    def _add_letter_content(self, doc, content):
        """Ajouter le contenu de la lettre avec le style approprié"""
        paragraphs = content.split('\n\n')
        
        for paragraph in paragraphs:
            if paragraph.strip():
                p = self._add_paragraph_with_style(doc, paragraph)
                p.paragraph_format.first_line_indent = Mm(10)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_paragraph_with_style(self, doc, text, alignment=None, bold=False):
        """Ajouter un paragraphe avec le style par défaut"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.default_font
        run.font.size = Pt(self.default_font_size)
        run.bold = bold
        
        # Style du paragraphe
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        
        if alignment:
            p.alignment = alignment
            
        return p

class StyleManager:
    """Gestionnaire de styles pour le contenu de la lettre"""
    
    def __init__(self):
        self.default_styles = {
            "font_family": "Times New Roman",
            "font_size": "11.5pt",
            "line_height": "1.15",
            "text_align": "left",
            "margin_top": "0",
            "margin_bottom": "0",
            "first_line_indent": "10mm"
        }
        
        self.text_styles = {
            "normal": {},
            "bold": {"font-weight": "bold"},
            "italic": {"font-style": "italic"},
            "underline": {"text-decoration": "underline"},
            "header": {
                "font-weight": "bold",
                "font-size": "14pt",
                "margin-bottom": "1em"
            },
            "signature": {
                "text-align": "right",
                "margin-top": "2em"
            },
            "bullet": {
                "list-style-type": "disc",
                "margin-left": "20mm"
            },
            "numbered": {
                "list-style-type": "decimal",
                "margin-left": "20mm"
            }
        }
        
        self.paragraph_styles = {
            "body": {
                "text-align": "justify",
                "text-indent": "10mm",
                "margin-bottom": "1em"
            },
            "address": {
                "text-align": "left",
                "margin-bottom": "0.5em"
            },
            "date": {
                "text-align": "right",
                "margin-bottom": "2em"
            },
            "greeting": {
                "text-align": "left",
                "margin-bottom": "1.5em"
            },
            "closing": {
                "text-align": "left",
                "margin-top": "1.5em",
                "margin-bottom": "1.5em"
            },
            "list-item": {
                "text-align": "left",
                "margin-bottom": "0.5em",
                "text-indent": "0"
            },
            "table": {
                "width": "100%",
                "border-collapse": "collapse",
                "margin-bottom": "1em"
            },
            "table-cell": {
                "border": "1px solid #ddd",
                "padding": "8px",
                "text-align": "left"
            },
            "header-footer": {
                "font-size": "9pt",
                "text-align": "center",
                "margin-top": "5mm",
                "margin-bottom": "5mm",
                "color": "#666"
            }
        }
        
        # Styles pour l'impression
        self.print_styles = {
            "@page": {
                "size": "A4",
                "margin": "25mm 20mm"
            },
            "@media print": {
                "body": {
                    "width": "210mm",
                    "height": "297mm",
                    "margin": "0 auto"
                }
            }
        }
    
    def get_style(self, style_name, style_type="text"):
        """Récupérer un style par son nom"""
        if style_type == "text":
            return self.text_styles.get(style_name, {})
        elif style_type == "paragraph":
            return self.paragraph_styles.get(style_name, {})
        return {}
    
    def apply_styles(self, content, styles):
        """Appliquer des styles au contenu"""
        style_str = "; ".join([f"{k}: {v}" for k, v in styles.items()])
        return f'<div style="{style_str}">{content}</div>'
    
    def create_css(self):
        """Générer le CSS pour tous les styles"""
        css = []
        
        # Styles par défaut
        css.append(".letter-content {")
        for prop, value in self.default_styles.items():
            css.append(f"    {prop.replace('_', '-')}: {value};")
        css.append("}")
        
        # Styles de texte
        for name, styles in self.text_styles.items():
            css.append(f".text-{name} {{")
            for prop, value in styles.items():
                css.append(f"    {prop}: {value};")
            css.append("}")
        
        # Styles de paragraphe
        for name, styles in self.paragraph_styles.items():
            css.append(f".paragraph-{name} {{")
            for prop, value in styles.items():
                css.append(f"    {prop}: {value};")
            css.append("}")
        
        # Styles d'impression
        for selector, styles in self.print_styles.items():
            css.append(f"{selector} {{")
            for prop, value in styles.items():
                css.append(f"    {prop}: {value};")
            css.append("}")
        
        return "\n".join(css)

class LetterFormatter:
    """Classe pour formater la lettre de motivation"""
    
    def __init__(self):
        self.style_manager = StyleManager()
    
    def format_letter(self, data):
        """Formater une lettre complète"""
        sections = []
        
        # En-tête
        if data.get('header'):
            sections.append(self._format_header(data['header']))
        
        # Informations de l'expéditeur
        sections.append(self._format_sender_info(data))
        
        # Informations du destinataire
        sections.append(self._format_recipient_info(data))
        
        # Date et lieu
        sections.append(self._format_date_section(data))
        
        # Objet
        sections.append(self._format_subject(data))
        
        # Corps de la lettre
        sections.append(self._format_body(data))
        
        # Signature
        sections.append(self._format_signature(data))
        
        # Pied de page
        if data.get('footer'):
            sections.append(self._format_footer(data['footer']))
        
        # Assembler toutes les sections
        letter_content = "\n".join(sections)
        
        # Appliquer les styles globaux
        return self.style_manager.apply_styles(
            letter_content,
            self.style_manager.default_styles
        )
    
    def _format_sender_info(self, data):
        """Formater les informations de l'expéditeur"""
        sender_info = [
            data['full_name'],
            data['address'],
            f"{data['postal_code']} {data['city']}",
            f"Tél : {data['phone']}",
            f"Email : {data['email']}"
        ]
        
        return self.style_manager.apply_styles(
            "\n".join([f"<p>{line}</p>" for line in sender_info]),
            self.style_manager.get_style("address", "paragraph")
        )
    
    def _format_recipient_info(self, data):
        """Formater les informations du destinataire"""
        recipient_info = [
            data['company'],
            data['company_address'],
            f"{data['company_postal_code']} {data['company_city']}"
        ]
        
        return self.style_manager.apply_styles(
            "\n".join([f"<p>{line}</p>" for line in recipient_info]),
            self.style_manager.get_style("address", "paragraph")
        )
    
    def _format_date_section(self, data):
        """Formater la section date et lieu"""
        date_text = f"{data['city']}, le {data['date']}"
        
        return self.style_manager.apply_styles(
            f"<p>{date_text}</p>",
            self.style_manager.get_style("date", "paragraph")
        )
    
    def _format_subject(self, data):
        """Formater la ligne d'objet"""
        subject_text = f"Objet : {data['subject']}"
        
        return self.style_manager.apply_styles(
            f"<p>{subject_text}</p>",
            {**self.style_manager.get_style("bold", "text"),
             **self.style_manager.get_style("body", "paragraph")}
        )
    
    def _format_body(self, data):
        """Formater le corps de la lettre"""
        # Séparer le contenu en paragraphes
        paragraphs = data['content'].split('\n\n')
        
        formatted_paragraphs = []
        for paragraph in paragraphs:
            # Détecter si c'est une liste
            if paragraph.startswith('- '):
                items = paragraph.split('\n- ')
                items[0] = items[0][2:]  # Enlever le "- " du premier item
                formatted_list = "\n".join([
                    f'<li>{item}</li>' for item in items
                ])
                formatted_paragraphs.append(
                    self.style_manager.apply_styles(
                        f'<ul>{formatted_list}</ul>',
                        self.style_manager.get_style("bullet", "text")
                    )
                )
            else:
                formatted_paragraphs.append(
                    self.style_manager.apply_styles(
                        f'<p>{paragraph}</p>',
                        self.style_manager.get_style("body", "paragraph")
                    )
                )
        
        return "\n".join(formatted_paragraphs)
    
    def _format_signature(self, data):
        """Formater la signature"""
        return self.style_manager.apply_styles(
            f"<p>{data['full_name']}</p>",
            self.style_manager.get_style("signature", "text")
        )
    
    def _format_header(self, header_text):
        """Formater l'en-tête"""
        return self.style_manager.apply_styles(
            f"<div>{header_text}</div>",
            self.style_manager.get_style("header-footer", "paragraph")
        )
    
    def _format_footer(self, footer_text):
        """Formater le pied de page"""
        return self.style_manager.apply_styles(
            f"<div>{footer_text}</div>",
            self.style_manager.get_style("header-footer", "paragraph")
        )

# Initialiser le formateur de lettre
letter_formatter = LetterFormatter()

generator = LetterGenerator()
template_manager = TemplateManager(generator.save_dir)

class DocumentManager:
    """Gestionnaire de documents pour la création et l'export des lettres"""
    
    def __init__(self):
        self.default_font = "Times New Roman"
        self.default_font_size = 11.5
        self.default_line_spacing = 1.15
        self.default_margins = {
            "top": 25,    # mm
            "bottom": 25,
            "left": 20,
            "right": 20
        }
        self.page_size = {
            "width": 210,  # A4 en mm
            "height": 297
        }
        
        # Dossier pour les fichiers temporaires
        self.temp_dir = os.path.join(tempfile.gettempdir(), "lettre_motivation_ai")
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def create_document(self, data, output_format="docx"):
        """Créer un document avec les données fournies"""
        try:
            # Créer un document Word temporaire
            doc = Document()
            
            # Configurer les marges
            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(self.default_margins["top"])
                section.bottom_margin = Mm(self.default_margins["bottom"])
                section.left_margin = Mm(self.default_margins["left"])
                section.right_margin = Mm(self.default_margins["right"])
            
            # En-tête personnalisé
            if data.get('header'):
                self._add_header(doc, data['header'])
            
            # Informations de l'expéditeur
            self._add_sender_info(doc, data)
            
            # Espace
            doc.add_paragraph()
            
            # Informations du destinataire
            self._add_recipient_info(doc, data)
            
            # Espace
            doc.add_paragraph()
            
            # Date et lieu
            self._add_date_location(doc, data)
            
            # Espace
            doc.add_paragraph()
            
            # Objet
            self._add_subject(doc, data)
            
            # Espace
            doc.add_paragraph()
            
            # Formule de politesse début
            self._add_paragraph_with_style(doc, "Madame, Monsieur,")
            
            # Espace
            doc.add_paragraph()
            
            # Contenu principal
            self._add_main_content(doc, data)
            
            # Formule de politesse fin
            self._add_closing(doc)
            
            # Signature
            self._add_signature(doc, data)
            
            # Pied de page personnalisé
            if data.get('footer'):
                self._add_footer(doc, data['footer'])
            
            # Créer un fichier temporaire
            temp_docx = os.path.join(
                self.temp_dir,
                f"lettre_motivation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            doc.save(temp_docx)
            
            # Convertir en PDF si demandé
            if output_format == "pdf":
                temp_pdf = temp_docx.replace('.docx', '.pdf')
                convert(temp_docx, temp_pdf)
                
                # Lire le fichier PDF
                with open(temp_pdf, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Nettoyer
                os.unlink(temp_pdf)
                os.unlink(temp_docx)
                
                return pdf_content
            
            # Sinon, renvoyer le fichier Word
            else:
                with open(temp_docx, 'rb') as docx_file:
                    docx_content = docx_file.read()
                
                # Nettoyer
                os.unlink(temp_docx)
                
                return docx_content
            
        except Exception as e:
            raise Exception(f"Erreur lors de la création du document : {str(e)}")
    
    def _add_header(self, doc, header_text):
        """Ajouter un en-tête personnalisé"""
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = header_text
        header_para.style = doc.styles['Header']
        
        # Style de l'en-tête
        font = header_para.style.font
        font.name = self.default_font
        font.size = Pt(9)
        font.color.rgb = RGBColor(102, 102, 102)
    
    def _add_footer(self, doc, footer_text):
        """Ajouter un pied de page personnalisé"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = footer_text
        footer_para.style = doc.styles['Footer']
        
        # Style du pied de page
        font = footer_para.style.font
        font.name = self.default_font
        font.size = Pt(9)
        font.color.rgb = RGBColor(102, 102, 102)
    
    def _add_sender_info(self, doc, data):
        """Ajouter les informations de l'expéditeur"""
        for info in [
            data['full_name'],
            data['address'],
            f"{data['postal_code']} {data['city']}",
            f"Tél : {data['phone']}",
            f"Email : {data['email']}"
        ]:
            self._add_paragraph_with_style(doc, info)
    
    def _add_recipient_info(self, doc, data):
        """Ajouter les informations du destinataire"""
        for info in [
            data['company'],
            data['company_address'],
            f"{data['company_postal_code']} {data['company_city']}"
        ]:
            self._add_paragraph_with_style(doc, info)
    
    def _add_date_location(self, doc, data):
        """Ajouter la date et le lieu"""
        p = self._add_paragraph_with_style(
            doc,
            f"{data['city']}, le {data['date']}",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )
    
    def _add_subject(self, doc, data):
        """Ajouter l'objet"""
        p = doc.add_paragraph()
        run = p.add_run("Objet : ")
        self._apply_run_style(run, bold=True)
        
        run = p.add_run(data['subject'])
        self._apply_run_style(run)
    
    def _add_main_content(self, doc, data):
        """Ajouter le contenu principal"""
        paragraphs = data['content'].split('\n')
        
        for para in paragraphs:
            if para.strip():
                # Détecter si c'est une liste
                if para.startswith('- '):
                    items = para.split('\n- ')
                    items[0] = items[0][2:]  # Enlever le "- " du premier item
                    
                    for item in items:
                        p = self._add_paragraph_with_style(
                            doc,
                            item,
                            first_line_indent=None,
                            left_indent=Mm(10)
                        )
                        p.style = 'List Bullet'
                else:
                    self._add_paragraph_with_style(
                        doc,
                        para,
                        first_line_indent=Mm(10)
                    )
    
    def _add_closing(self, doc):
        """Ajouter la formule de politesse finale"""
        self._add_paragraph_with_style(
            doc,
            "Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées."
        )
    
    def _add_signature(self, doc, data):
        """Ajouter la signature"""
        self._add_paragraph_with_style(
            doc,
            data['full_name'],
            alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )
    
    def _add_paragraph_with_style(self, doc, text, alignment=None, bold=False,
                                first_line_indent=None, left_indent=None):
        """Ajouter un paragraphe avec style"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        self._apply_run_style(run, bold)
        
        # Format du paragraphe
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = self.default_line_spacing
        
        if alignment:
            p.alignment = alignment
        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent
        if left_indent is not None:
            pf.left_indent = left_indent
        
        return p
    
    def _apply_run_style(self, run, bold=False):
        """Appliquer le style à un run"""
        run.font.name = self.default_font
        run.font.size = Pt(self.default_font_size)
        run.bold = bold

# Initialiser le gestionnaire de documents
document_manager = DocumentManager()

def process_text_formatting(text):
    """Traite le texte formaté et retourne le texte avec le formatage HTML."""
    import re
    
    # Traiter les balises de style
    text = re.sub(r'<bold>(.*?)</bold>', r'<strong>\1</strong>', text)
    text = re.sub(r'<italic>(.*?)</italic>', r'<em>\1</em>', text)
    text = re.sub(r'<underline>(.*?)</underline>', r'<u>\1</u>', text)
    
    # Traiter les alignements
    text = re.sub(r'<align="(.*?)">(.*?)</align>', r'<div style="text-align: \1">\2</div>', text)
    
    # Traiter l'interligne
    text = re.sub(r'<spacing="(.*?)">(.*?)</spacing>', r'<div style="line-height: \1">\2</div>', text)
    
    return text

def process_long_text(text, max_length=1000):
    """Traite un texte long en le divisant en sections."""
    sections = []
    words = text.split()
    current_section = []
    current_length = 0
    
    for word in words:
        if current_length + len(word) + 1 <= max_length:
            current_section.append(word)
            current_length += len(word) + 1
        else:
            sections.append(" ".join(current_section))
            current_section = [word]
            current_length = len(word)
    
    if current_section:
        sections.append(" ".join(current_section))
    
    return sections

def create_word_document(content, filename):
    """Crée un document Word avec le contenu formaté."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    # Créer un nouveau document
    doc = Document()
    
    # Configurer les marges (1 pouce = 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Fonction pour appliquer le style au paragraphe
    def apply_paragraph_style(paragraph, style_dict):
        if 'alignment' in style_dict:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            paragraph.alignment = align_map.get(style_dict['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
        
        if 'line_spacing' in style_dict:
            paragraph.paragraph_format.line_spacing = float(style_dict['line_spacing'])
    
    # Fonction pour appliquer le style au texte
    def apply_run_style(run, style_dict):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        if style_dict.get('bold'):
            run.font.bold = True
        if style_dict.get('italic'):
            run.font.italic = True
        if style_dict.get('underline'):
            run.font.underline = True
    
    # Parser le contenu HTML et créer les paragraphes
    paragraphs = content.split('\n')
    for p_text in paragraphs:
        if not p_text.strip():
            continue
        
        # Créer un nouveau paragraphe
        p = doc.add_paragraph()
        
        # Extraire les styles du paragraphe
        style_dict = {}
        
        # Vérifier l'alignement
        align_match = re.search(r'text-align:\s*(\w+)', p_text)
        if align_match:
            style_dict['alignment'] = align_match.group(1)
        
        # Vérifier l'interligne
        spacing_match = re.search(r'line-height:\s*([\d.]+)', p_text)
        if spacing_match:
            style_dict['line_spacing'] = float(spacing_match.group(1))
        
        # Appliquer les styles au paragraphe
        apply_paragraph_style(p, style_dict)
        
        # Traiter le texte avec les balises de style
        text = re.sub(r'<div[^>]*>', '', p_text)
        text = text.replace('</div>', '')
        
        # Diviser le texte en segments basés sur le formatage
        segments = re.split(r'(</?(?:strong|em|u)>)', text)
        current_styles = {'bold': False, 'italic': False, 'underline': False}
        
        for segment in segments:
            if segment in ('<strong>', '<em>', '<u>'):
                if segment == '<strong>': current_styles['bold'] = True
                elif segment == '<em>': current_styles['italic'] = True
                elif segment == '<u>': current_styles['underline'] = True
            elif segment in ('</strong>', '</em>', '</u>'):
                if segment == '</strong>': current_styles['bold'] = False
                elif segment == '</em>': current_styles['italic'] = False
                elif segment == '</u>': current_styles['underline'] = False
            else:
                if segment.strip():
                    run = p.add_run(segment)
                    apply_run_style(run, current_styles)
    
    # Sauvegarder le document
    doc.save(filename)
    return filename

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/export', methods=['POST'])
def export_letter():
    """
    Exporte la lettre au format Word ou PDF
    """
    try:
        data = request.get_json()
        
        # Valider les données
        required_fields = ['full_name', 'address', 'postal_code', 'city', 'phone', 'email',
                         'company', 'company_address', 'company_postal_code', 'company_city',
                         'subject', 'content', 'format']
        
        for field in required_fields:
            if not data.get(field):
                return jsonify({
                    'success': False,
                    'error': f'Le champ {field} est requis'
                }), 400
        
        if data['format'] not in ['docx', 'pdf']:
            return jsonify({
                'success': False,
                'error': 'Format non supporté'
            }), 400
        
        # Créer le document
        file_content = document_manager.create_document(data, data['format'])
        
        # Déterminer le type MIME
        mime_type = ('application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    if data['format'] == 'docx' else 'application/pdf')
        
        # Renvoyer le fichier
        return send_file(
            io.BytesIO(file_content),
            mimetype=mime_type,
            as_attachment=True,
            download_name=f"lettre_motivation_{datetime.now().strftime('%Y-%m-%d')}.{data['format']}"
        )
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/health')
def health_check():
    """Route de vérification de santé pour Render"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat()
    })

if __name__ == '__main__':
    app.run(debug=True)
