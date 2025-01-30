import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from docx2pdf import convert
import threading
from pathlib import Path

class LetterGeneratorApp:
    def __init__(self):
        self.window = ctk.CTk()
        self.window.title("Générateur de Lettre de Motivation")
        self.window.geometry("800x900")
        
        # Configuration des styles
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # Variables pour les champs de saisie
        self.company_var = ctk.StringVar()
        self.position_var = ctk.StringVar()
        self.duration_var = ctk.StringVar()
        self.start_date_var = ctk.StringVar()
        self.today_date_var = ctk.StringVar(value=self.get_current_date())
        self.output_folder_var = ctk.StringVar(value=str(Path.home() / "Documents"))
        self.filename_var = ctk.StringVar(value="lettre_motivation")
        
        # Charger les templates
        self.templates = self.load_templates()
        self.selected_template = None
        
        self.create_widgets()
        self.load_last_session()
    
    def create_widgets(self):
        # Frame principal
        main_frame = ctk.CTkFrame(self.window)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Section des entrées
        self.create_input_section(main_frame)
        
        # Section des templates
        self.create_template_section(main_frame)
        
        # Zone de texte personnalisé
        self.create_custom_text_section(main_frame)
        
        # Boutons d'action
        self.create_action_buttons(main_frame)
    
    def create_input_section(self, parent):
        input_frame = ctk.CTkFrame(parent)
        input_frame.pack(fill="x", padx=5, pady=5)
        
        # Entreprise
        ctk.CTkLabel(input_frame, text="Nom de l'entreprise:").pack(anchor="w", padx=5, pady=2)
        company_entry = ctk.CTkEntry(input_frame, textvariable=self.company_var)
        company_entry.pack(fill="x", padx=5, pady=2)
        
        # Poste
        ctk.CTkLabel(input_frame, text="Intitulé du poste:").pack(anchor="w", padx=5, pady=2)
        position_entry = ctk.CTkEntry(input_frame, textvariable=self.position_var)
        position_entry.pack(fill="x", padx=5, pady=2)
        
        # Durée
        ctk.CTkLabel(input_frame, text="Durée (optionnel):").pack(anchor="w", padx=5, pady=2)
        duration_entry = ctk.CTkEntry(input_frame, textvariable=self.duration_var)
        duration_entry.pack(fill="x", padx=5, pady=2)
        
        # Date de début
        ctk.CTkLabel(input_frame, text="Date de début (optionnel):").pack(anchor="w", padx=5, pady=2)
        start_date_entry = ctk.CTkEntry(input_frame, textvariable=self.start_date_var)
        start_date_entry.pack(fill="x", padx=5, pady=2)
        
        # Date du jour
        ctk.CTkLabel(input_frame, text="Date du jour:").pack(anchor="w", padx=5, pady=2)
        today_date_entry = ctk.CTkEntry(input_frame, textvariable=self.today_date_var)
        today_date_entry.pack(fill="x", padx=5, pady=2)
        
        # Dossier de sortie
        output_frame = ctk.CTkFrame(input_frame)
        output_frame.pack(fill="x", padx=5, pady=2)
        
        ctk.CTkLabel(output_frame, text="Dossier de sortie:").pack(side="left", padx=5)
        ctk.CTkEntry(output_frame, textvariable=self.output_folder_var).pack(side="left", fill="x", expand=True, padx=5)
        ctk.CTkButton(output_frame, text="Choisir", command=self.choose_output_folder).pack(side="right", padx=5)
        
        # Nom du fichier
        filename_frame = ctk.CTkFrame(input_frame)
        filename_frame.pack(fill="x", padx=5, pady=2)
        
        ctk.CTkLabel(filename_frame, text="Nom du fichier:").pack(side="left", padx=5)
        ctk.CTkEntry(filename_frame, textvariable=self.filename_var).pack(side="left", fill="x", expand=True, padx=5)
    
    def create_template_section(self, parent):
        template_frame = ctk.CTkFrame(parent)
        template_frame.pack(fill="x", padx=5, pady=5)
        
        # Liste déroulante des templates
        template_list_frame = ctk.CTkFrame(template_frame)
        template_list_frame.pack(fill="x", padx=5, pady=2)
        
        ctk.CTkLabel(template_list_frame, text="Template:").pack(side="left", padx=5)
        self.template_menu = ctk.CTkOptionMenu(
            template_list_frame,
            values=list(self.templates.keys()),
            command=self.on_template_selected
        )
        self.template_menu.pack(side="left", fill="x", expand=True, padx=5)
        
        # Boutons de gestion des templates
        button_frame = ctk.CTkFrame(template_frame)
        button_frame.pack(fill="x", padx=5, pady=2)
        
        ctk.CTkButton(button_frame, text="Nouveau", command=self.new_template).pack(side="left", padx=2)
        ctk.CTkButton(button_frame, text="Renommer", command=self.rename_template).pack(side="left", padx=2)
        ctk.CTkButton(button_frame, text="Supprimer", command=self.delete_template).pack(side="left", padx=2)
    
    def create_custom_text_section(self, parent):
        text_frame = ctk.CTkFrame(parent)
        text_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        ctk.CTkLabel(text_frame, text="Texte personnalisé:").pack(anchor="w", padx=5, pady=2)
        
        self.custom_text = ctk.CTkTextbox(text_frame, height=200)
        self.custom_text.pack(fill="both", expand=True, padx=5, pady=2)
    
    def create_action_buttons(self, parent):
        button_frame = ctk.CTkFrame(parent)
        button_frame.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkButton(button_frame, text="Générer Word", command=self.generate_word).pack(side="left", padx=5)
        ctk.CTkButton(button_frame, text="Générer PDF", command=self.generate_pdf).pack(side="left", padx=5)
        ctk.CTkButton(button_frame, text="Sauvegarder Template", command=self.save_template).pack(side="left", padx=5)
    
    def choose_output_folder(self):
        folder = filedialog.askdirectory(initialdir=self.output_folder_var.get())
        if folder:
            self.output_folder_var.set(folder)
    
    def get_current_date(self):
        return datetime.now().strftime("%d %B %Y")
    
    def load_templates(self):
        try:
            with open("templates.json", "r", encoding="utf-8") as f:
                return json.load(f)
        except FileNotFoundError:
            return {
                "Template par défaut": "Je suis très intéressé(e) par ce poste car il correspond parfaitement à mes compétences et à mes aspirations professionnelles."
            }
    
    def save_templates(self):
        with open("templates.json", "w", encoding="utf-8") as f:
            json.dump(self.templates, f, ensure_ascii=False, indent=2)
    
    def on_template_selected(self, template_name):
        self.selected_template = template_name
        if template_name in self.templates:
            self.custom_text.delete("1.0", "end")
            self.custom_text.insert("1.0", self.templates[template_name])
    
    def new_template(self):
        name = ctk.CTkInputDialog(
            text="Entrez le nom du nouveau template:",
            title="Nouveau Template"
        ).get_input()
        
        if name and name not in self.templates:
            self.templates[name] = ""
            self.save_templates()
            self.update_template_menu()
            self.template_menu.set(name)
            self.selected_template = name
    
    def rename_template(self):
        if not self.selected_template:
            messagebox.showwarning("Attention", "Veuillez sélectionner un template à renommer.")
            return
        
        new_name = ctk.CTkInputDialog(
            text="Entrez le nouveau nom du template:",
            title="Renommer Template"
        ).get_input()
        
        if new_name and new_name != self.selected_template:
            self.templates[new_name] = self.templates[self.selected_template]
            del self.templates[self.selected_template]
            self.save_templates()
            self.update_template_menu()
            self.template_menu.set(new_name)
            self.selected_template = new_name
    
    def delete_template(self):
        if not self.selected_template:
            messagebox.showwarning("Attention", "Veuillez sélectionner un template à supprimer.")
            return
        
        if messagebox.askyesno("Confirmation", f"Voulez-vous vraiment supprimer le template '{self.selected_template}' ?"):
            del self.templates[self.selected_template]
            self.save_templates()
            self.update_template_menu()
            if self.templates:
                first_template = list(self.templates.keys())[0]
                self.template_menu.set(first_template)
                self.selected_template = first_template
                self.custom_text.delete("1.0", "end")
                self.custom_text.insert("1.0", self.templates[first_template])
            else:
                self.selected_template = None
                self.custom_text.delete("1.0", "end")
    
    def update_template_menu(self):
        self.template_menu.configure(values=list(self.templates.keys()))
    
    def save_template(self):
        if not self.selected_template:
            messagebox.showwarning("Attention", "Veuillez sélectionner un template à mettre à jour.")
            return
        
        self.templates[self.selected_template] = self.custom_text.get("1.0", "end-1c")
        self.save_templates()
        messagebox.showinfo("Succès", "Template sauvegardé avec succès!")
    
    def create_letter(self, output_path):
        doc = Document()
        
        # Style par défaut
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(self.today_date_var.get())
        
        # Saut de ligne
        doc.add_paragraph()
        
        # Formule de politesse
        doc.add_paragraph('Madame, Monsieur,')
        
        # Corps de la lettre
        body = doc.add_paragraph()
        body.add_run(f'Je me permets de vous adresser ma candidature pour le poste de {self.position_var.get()} au sein de votre entreprise {self.company_var.get()}')
        
        if self.duration_var.get():
            body.add_run(f' pour une durée de {self.duration_var.get()}')
        
        if self.start_date_var.get():
            body.add_run(f', à partir du {self.start_date_var.get()}')
        
        body.add_run('.')
        
        # Texte personnalisé
        custom_text = self.custom_text.get("1.0", "end-1c")
        if custom_text:
            doc.add_paragraph()
            doc.add_paragraph(custom_text)
        
        # Formule de fin
        doc.add_paragraph()
        doc.add_paragraph('Je me tiens à votre disposition pour un entretien et vous prie d\'agréer, Madame, Monsieur, l\'expression de mes salutations distinguées.')
        
        # Sauvegarder
        doc.save(output_path)
    
    def generate_word(self):
        try:
            output_path = os.path.join(
                self.output_folder_var.get(),
                f"{self.filename_var.get()}.docx"
            )
            
            self.create_letter(output_path)
            self.save_session()
            messagebox.showinfo("Succès", f"Document Word généré avec succès!\nChemin: {output_path}")
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Une erreur est survenue: {str(e)}")
    
    def generate_pdf(self):
        try:
            # Créer d'abord le document Word
            word_path = os.path.join(
                self.output_folder_var.get(),
                f"{self.filename_var.get()}.docx"
            )
            self.create_letter(word_path)
            
            # Convertir en PDF
            pdf_path = os.path.join(
                self.output_folder_var.get(),
                f"{self.filename_var.get()}.pdf"
            )
            
            # Utiliser un thread séparé pour la conversion
            def convert_to_pdf():
                try:
                    convert(word_path, pdf_path)
                    self.window.after(0, lambda: messagebox.showinfo("Succès", f"Document PDF généré avec succès!\nChemin: {pdf_path}"))
                except Exception as e:
                    self.window.after(0, lambda: messagebox.showerror("Erreur", f"Erreur lors de la conversion en PDF: {str(e)}"))
            
            thread = threading.Thread(target=convert_to_pdf)
            thread.start()
            
            self.save_session()
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Une erreur est survenue: {str(e)}")
    
    def save_session(self):
        session = {
            "company": self.company_var.get(),
            "position": self.position_var.get(),
            "duration": self.duration_var.get(),
            "start_date": self.start_date_var.get(),
            "output_folder": self.output_folder_var.get(),
            "filename": self.filename_var.get(),
            "selected_template": self.selected_template,
            "custom_text": self.custom_text.get("1.0", "end-1c")
        }
        
        with open("last_session.json", "w", encoding="utf-8") as f:
            json.dump(session, f, ensure_ascii=False, indent=2)
    
    def load_last_session(self):
        try:
            with open("last_session.json", "r", encoding="utf-8") as f:
                session = json.load(f)
                
                self.company_var.set(session.get("company", ""))
                self.position_var.set(session.get("position", ""))
                self.duration_var.set(session.get("duration", ""))
                self.start_date_var.set(session.get("start_date", ""))
                self.output_folder_var.set(session.get("output_folder", str(Path.home() / "Documents")))
                self.filename_var.set(session.get("filename", "lettre_motivation"))
                
                if "selected_template" in session and session["selected_template"] in self.templates:
                    self.template_menu.set(session["selected_template"])
                    self.selected_template = session["selected_template"]
                
                if "custom_text" in session:
                    self.custom_text.delete("1.0", "end")
                    self.custom_text.insert("1.0", session["custom_text"])
        
        except FileNotFoundError:
            pass
    
    def run(self):
        self.window.mainloop()

if __name__ == "__main__":
    app = LetterGeneratorApp()
    app.run()
